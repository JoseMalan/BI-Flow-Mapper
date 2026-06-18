from __future__ import annotations

import json
import re
import socket
import sys
import tempfile
import threading
import traceback
import webbrowser
import zipfile
from datetime import datetime
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlparse
from urllib.request import urlopen

from connector_catalog import POWER_QUERY_CONNECTORS

# Heavy libraries are loaded only when the user analyzes a PBIX or exports Word.
# Keeping them out of the desktop launch path makes the WebView appear sooner.
PBIXRay = None
PBIXRAY_IMPORT_ERROR = ""
_PBIXRAY_IMPORT_ATTEMPTED = False

Document = None
DOCX_IMPORT_ERROR = ""
_DOCX_IMPORT_ATTEMPTED = False

PIL_AVAILABLE = False
_PIL_IMPORT_ATTEMPTED = False
_PILImage = _PILDraw = _PILFont = None

CAIROSVG_AVAILABLE = False
_CAIROSVG_IMPORT_ATTEMPTED = False
_cairosvg = None


def load_pbixray() -> bool:
    global PBIXRay, PBIXRAY_IMPORT_ERROR, _PBIXRAY_IMPORT_ATTEMPTED
    if not _PBIXRAY_IMPORT_ATTEMPTED:
        _PBIXRAY_IMPORT_ATTEMPTED = True
        try:
            from pbixray import PBIXRay as _PBIXRay
            PBIXRay = _PBIXRay
        except Exception as import_error:
            PBIXRAY_IMPORT_ERROR = str(import_error)
    return PBIXRay is not None


def load_docx() -> bool:
    global Document, DOCX_IMPORT_ERROR, _DOCX_IMPORT_ATTEMPTED
    global WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH
    global OxmlElement, qn, Inches, Pt, RGBColor
    if not _DOCX_IMPORT_ATTEMPTED:
        _DOCX_IMPORT_ATTEMPTED = True
        try:
            from docx import Document as _Document
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as _WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement as _OxmlElement
            from docx.oxml.ns import qn as _qn
            from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor
            Document = _Document
            WD_CELL_VERTICAL_ALIGNMENT = _WD_CELL_VERTICAL_ALIGNMENT
            WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
            WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
            OxmlElement = _OxmlElement
            qn = _qn
            Inches, Pt, RGBColor = _Inches, _Pt, _RGBColor
        except Exception as import_error:
            DOCX_IMPORT_ERROR = str(import_error)
    return Document is not None


def load_pillow() -> bool:
    global PIL_AVAILABLE, _PIL_IMPORT_ATTEMPTED
    global _PILImage, _PILDraw, _PILFont
    if not _PIL_IMPORT_ATTEMPTED:
        _PIL_IMPORT_ATTEMPTED = True
        try:
            from PIL import Image, ImageDraw, ImageFont
            _PILImage, _PILDraw, _PILFont = Image, ImageDraw, ImageFont
            PIL_AVAILABLE = True
        except Exception:
            PIL_AVAILABLE = False
    return PIL_AVAILABLE


def load_cairosvg() -> bool:
    global CAIROSVG_AVAILABLE, _CAIROSVG_IMPORT_ATTEMPTED, _cairosvg
    if not _CAIROSVG_IMPORT_ATTEMPTED:
        _CAIROSVG_IMPORT_ATTEMPTED = True
        try:
            import cairosvg
            _cairosvg = cairosvg
            CAIROSVG_AVAILABLE = True
        except Exception:
            CAIROSVG_AVAILABLE = False
    return CAIROSVG_AVAILABLE

# Pasta dos assets empacotados (index.html, app.js, styles.css...)
if getattr(sys, 'frozen', False):
    ASSETS_DIR = Path(sys._MEIPASS)          # onde o PyInstaller extrai os arquivos
    RUNTIME_DIR = Path(sys.executable).parent # onde o .exe está (para o port file)
else:
    ASSETS_DIR = Path(__file__).resolve().parent
    RUNTIME_DIR = ASSETS_DIR

ROOT = ASSETS_DIR
PORT_FILE = RUNTIME_DIR / "bi-flow-mapper.port"
HOST = "127.0.0.1"
DEFAULT_PORT = 4173

CONNECTORS = POWER_QUERY_CONNECTORS
CANONICAL_CONNECTOR_NAMES = {
    "Azure Database for PostgreSQL": "PostgreSQL",
    "PostgreSQL database": "PostgreSQL",
    "Azure SQL Database": "SQL Server",
    "Azure Synapse Analytics (SQL Data Warehouse)": "SQL Server",
    "SQL Server database": "SQL Server",
    "Excel Workbook": "Excel",
    "Text/CSV": "CSV",
    "JSON": "JSON",
    "SharePoint Folder": "SharePoint",
    "SharePoint Online List": "SharePoint",
    "OData feed": "OData",
    "Oracle database": "Oracle",
    "MySQL database": "MySQL",
    "IBM Db2 database": "IBM Db2",
    "Informix database": "Informix",
    "Sybase database": "Sybase",
    "Teradata database": "Teradata",
}
PREFERRED_PATTERN_CONNECTORS = {
    "Access.Database": "Access Database",
    "ActiveDirectory.Domains": "Active Directory",
    "AdobeAnalytics.Cubes": "Adobe Analytics",
    "AmazonRedshift.Database": "Amazon Redshift",
    "AnalysisServices.Database": "Analysis Services",
    "AzureDataExplorer.Contents": "Azure Data Explorer (Kusto)",
    "Kusto.Contents": "Azure Data Explorer (Kusto)",
    "AzureDataLakeStorage.Contents": "Azure Data Lake Storage Gen1",
    "AzureStorage.Blobs": "Azure Blob Storage",
    "AzureStorage.DataLake": "Azure Data Lake Storage Gen2",
    "AzureStorage.Tables": "Azure Table Storage",
    "CommonDataService.Database": "Dataverse",
    "Dataverse.Contents": "Dataverse",
    "Csv.Document": "CSV",
    "Databricks.Catalogs": "Databricks",
    "Databricks.Query": "Databricks",
    "Excel.Workbook": "Excel",
    "Folder.Files": "Folder",
    "Folder.Contents": "Folder",
    "GoogleAnalytics.Accounts": "Google Analytics",
    "GoogleBigQuery.Database": "Google BigQuery",
    "Hdfs.Files": "Hadoop File (HDFS)",
    "Hdfs.Contents": "Hadoop File (HDFS)",
    "DB2.Database": "IBM Db2 database",
    "Impala.Database": "Impala",
    "Informix.Database": "Informix database",
    "Json.Document": "JSON",
    "MySQL.Database": "MySQL database",
    "OData.Feed": "OData feed",
    "Odbc.DataSource": "ODBC",
    "Odbc.Query": "ODBC",
    "OleDb.DataSource": "OLE DB",
    "OleDb.Query": "OLE DB",
    "Oracle.Database": "Oracle database",
    "Parquet.Document": "Parquet",
    "Pdf.Tables": "PDF",
    "PostgreSQL.Database": "PostgreSQL",
    "PowerPlatform.Dataflows": "Power Platform Dataflows",
    "Salesforce.Data": "Salesforce Objects",
    "Salesforce.Objects": "Salesforce Objects",
    "Salesforce.Reports": "Salesforce Reports",
    "SapBusinessWarehouse.Cubes": "SAP Business Warehouse Application Server",
    "SharePoint.Files": "SharePoint",
    "SharePoint.Contents": "SharePoint",
    "SharePoint.Tables": "SharePoint Online List",
    "Snowflake.Databases": "Snowflake",
    "Sql.Database": "SQL Server",
    "Sql.Databases": "SQL Server",
    "Sybase.Database": "Sybase database",
    "Teradata.Database": "Teradata database",
    "Vertica.Database": "Vertica",
    "Web.Contents": "Web",
    "Web.BrowserContents": "Web",
    "Xml.Tables": "XML",
    "Xml.Document": "XML",
}


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def do_POST(self):
        parsed_url = urlparse(self.path)
        request_path = parsed_url.path
        if request_path not in {"/api/analyze", "/api/export-docx"}:
            self.send_json({"error": "Not found"}, status=404)
            return

        try:
            file_name, payload = self.read_upload()
            with tempfile.TemporaryDirectory(prefix="bi-flow-mapper-") as temp_dir:
                safe_name = re.sub(r"[^A-Za-z0-9_. -]+", "_", file_name or "upload.pbix")
                pbix_path = Path(temp_dir) / safe_name
                pbix_path.write_bytes(payload)

                if request_path == "/api/export-docx":
                    locale = self.headers.get("X-BIFM-Locale") or parse_qs(parsed_url.query).get("locale", ["pt-BR"])[0]
                    doc_locale = normalize_doc_locale(locale)
                    data = build_documentation_docx(pbix_path, file_name, doc_locale)
                    suffix = "documentacao" if doc_locale == "pt-BR" else "documentation"
                    docx_name = f"{Path(file_name or 'power-bi').stem}_{suffix}.docx"
                    self.send_binary(
                        data,
                        docx_name,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    return

                graph = analyze_pbix(pbix_path)
                graph["fileName"] = file_name
                self.send_json(graph)
        except Exception as error:
            traceback.print_exc()
            self.send_json({"error": str(error)}, status=500)

    def do_GET(self):
        if self.path == "/api/health":
            self.send_json({"ok": True, "app": "BI Flow Mapper"})
            return
        super().do_GET()

    def read_upload(self):
        content_type = self.headers.get("Content-Type", "")
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length)

        if "multipart/form-data" not in content_type:
            return "upload.pbix", body

        boundary_match = re.search(r"boundary=([^;]+)", content_type)
        if not boundary_match:
            raise ValueError("Upload multipart sem boundary.")

        boundary = ("--" + boundary_match.group(1).strip('"')).encode()
        for part in body.split(boundary):
            if b"\r\n\r\n" not in part:
                continue
            headers, payload = part.split(b"\r\n\r\n", 1)
            if b'name="pbix"' not in headers:
                continue
            payload = payload.rstrip(b"\r\n-")
            filename_match = re.search(rb'filename="([^"]+)"', headers)
            filename = "upload.pbix"
            if filename_match:
                filename = unquote(filename_match.group(1).decode("utf-8", errors="replace"))
            return filename, payload

        raise ValueError("Campo de upload 'pbix' nao encontrado.")

    def send_json(self, payload, status=200):
        data = json.dumps(payload, ensure_ascii=False, allow_nan=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_binary(self, data, file_name, content_type, status=200):
        safe_name = re.sub(r'["\r\n]+', "", file_name or "download.bin")
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Content-Disposition", f'attachment; filename="{safe_name}"')
        self.end_headers()
        self.wfile.write(data)


def analyze_pbix(path: Path):
    if not load_pbixray():
        raise RuntimeError(f"pbixray nao esta instalado ou nao carregou: {PBIXRAY_IMPORT_ERROR}")

    model = PBIXRay(str(path))
    diagnostics = []

    power_query = records_from(model, "power_query", diagnostics)
    datasource_records = records_from(model, "tmschema_datasources", diagnostics)
    measures = records_from(model, "dax_measures", diagnostics)
    calc_columns_raw = records_from(model, "dax_columns", diagnostics)
    relationships = records_from(model, "relationships", diagnostics)

    # Diagnóstico: exibe as colunas reais para facilitar debug
    try:
        raw_rels = getattr(model, "relationships", None)
        if raw_rels is not None and hasattr(raw_rels, "columns"):
            diagnostics.append(f"relationships columns: {list(raw_rels.columns)}")
            if len(raw_rels) > 0:
                diagnostics.append(f"relationships first row: {raw_rels.iloc[0].to_dict()}")
    except Exception as _e:
        diagnostics.append(f"relationships debug error: {_e}")
    schema = records_from(model, "schema", diagnostics)
    semantic_tables = records_from(model, "tmschema_tables", diagnostics)
    tables = list_from(model, "tables", diagnostics)
    visuals = extract_visuals_from_layout(path)

    connector_nodes = detect_connector_nodes(power_query, datasource_records)
    query_nodes = build_query_nodes(power_query, connector_nodes)
    table_nodes = build_table_nodes(tables, schema, semantic_tables, query_nodes)
    measure_nodes = build_measure_nodes(measures)
    calc_column_nodes = build_calc_column_nodes(calc_columns_raw)
    visual_nodes = build_visual_nodes(visuals)

    nodes = unique_by_id(connector_nodes + query_nodes + table_nodes + measure_nodes + calc_column_nodes + visual_nodes)
    edges = []
    valid_node_ids = {node["id"] for node in nodes}

    for query in query_nodes:
        query_text = query_search_text(query)
        matched_sources = [
            source for source in connector_nodes
            if source_matches_query(source, query_text)
        ]
        if not matched_sources and len(connector_nodes) == 1:
            matched_sources = connector_nodes
        for source in matched_sources:
            edges.append(edge(source["id"], query["id"], "uses connector"))

    table_by_label = {node["label"].lower(): node for node in table_nodes}
    for query in query_nodes:
        table = table_by_label.get(query["label"].lower())
        if table:
            edges.append(edge(query["id"], table["id"], "loads table"))

    # Relacionamentos entre tabelas do modelo NÃO entram no grafo do Mapa.
    # Eles são exibidos exclusivamente no painel "Relacionamentos" via structured_rels.

    for measure in measure_nodes:
        table_name = measure["meta"].get("table")
        if table_name and f"model:{slug(table_name)}" in valid_node_ids:
            edges.append(edge(f"model:{slug(table_name)}", measure["id"], "defines measure"))
        edges.extend(build_measure_dependency_edges(measure, measure_nodes, table_nodes))

    for calc_col in calc_column_nodes:
        table_name = calc_col["meta"].get("table")
        if table_name and f"model:{slug(table_name)}" in valid_node_ids:
            edges.append(edge(f"model:{slug(table_name)}", calc_col["id"], "defines calc column"))
        # Calc columns can also reference measures
        edges.extend(build_calc_column_dependency_edges(calc_col, measure_nodes, table_nodes))

    edges.extend(build_visual_edges(visual_nodes, measure_nodes, table_nodes, calc_column_nodes))

    if not nodes:
        nodes.append({
            "id": "file:pbix",
            "type": "source",
            "label": path.name,
            "icon": "PBX",
            "meta": {"expression": "\n".join(diagnostics) or "PBIXRay nao retornou metadados."},
        })

    diagnostics.append(f"Power Query rows: {len(power_query)}")
    diagnostics.append(f"Data source rows: {len(datasource_records)}")
    diagnostics.append(f"Tables: {len(table_nodes)}")
    diagnostics.append(f"Semantic table rows: {len(semantic_tables)}")
    diagnostics.append(f"Measures: {len(measure_nodes)}")
    diagnostics.append(f"Calculated columns: {len(calc_column_nodes)}")
    diagnostics.append(f"Relationships: {len(relationships)}")
    diagnostics.append(f"Visuals inferred from layout: {len(visual_nodes)}")

    structured_rels = build_structured_relationships(relationships)

    pages = extract_pages_from_layout(path)

    return {
        "nodes": nodes,
        "edges": unique_edges(edges),
        "warnings": diagnostics,
        "source": "pbixray",
        "relationships": structured_rels,
        "pages": pages,
    }


def svg_to_png_bytes(svg_string: str, scale: float = 2.0) -> bytes | None:
    """Convert SVG to PNG via cairosvg if available (best quality)."""
    if not svg_string:
        return None
    if load_cairosvg():
        try:
            return _cairosvg.svg2png(bytestring=svg_string.encode("utf-8"), scale=scale)
        except Exception:
            pass
    return None


def _hex(color: str) -> tuple:
    """Parse #RRGGBB or #RGB to (R,G,B)."""
    c = color.lstrip("#")
    if len(c) == 3:
        c = c[0]*2 + c[1]*2 + c[2]*2
    return (int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))


def _pil_png(img) -> bytes:
    from io import BytesIO as _B
    buf = _B()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_banner_png(text: str, color: str = "#0078D4", icon: str = "",
                    width: int = 1388, height: int = 56) -> bytes | None:
    """Draw a section banner using Pillow. No SVG needed."""
    if not PIL_AVAILABLE:
        return None
    try:
        img = _PILImage.new("RGBA", (width, height), (0, 0, 0, 0))
        draw = _PILDraw.Draw(img)
        r, g, b = _hex(color)
        # Tinted background
        bg = _PILImage.new("RGBA", (width, height), (r, g, b, 28))
        img = _PILImage.alpha_composite(img, bg)
        draw = _PILDraw.Draw(img)
        # Left accent bar
        draw.rectangle([0, 0, 8, height], fill=(r, g, b, 255))
        # Text (icon + label)
        label = (icon + "  " + text) if icon else text
        # Try to use a system font; fall back gracefully
        font = None
        if not font:
            try:
                font = _PILFont.truetype("arial.ttf", 24)
            except Exception:
                pass
        if not font:
            try:
                import os
                for p in ["C:/Windows/Fonts/segoeui.ttf",
                          "C:/Windows/Fonts/arial.ttf",
                          "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                          "/System/Library/Fonts/Helvetica.ttc"]:
                    if os.path.exists(p):
                        font = _PILFont.truetype(p, 24)
                        break
            except Exception:
                pass
        if not font:
            font = _PILFont.load_default()
        draw.text((24, height // 2 - 12), label, font=font, fill=(r, g, b, 255))
        return _pil_png(img)
    except Exception:
        return None


def make_erd_png(relationships: list, scale: int = 2) -> bytes | None:
    """Draw the ERD diagram using Pillow. No SVG needed."""
    if not PIL_AVAILABLE or not relationships:
        return None
    try:
        table_names = []
        seen_t: set = set()
        for rel in relationships:
            for t in [rel.get("fromTable",""), rel.get("toTable","")]:
                if t and t not in seen_t:
                    seen_t.add(t); table_names.append(t)
        table_names.sort()
        count  = len(table_names)
        NODE_W = 190; NODE_H = 48
        COLS   = min(count, max(2, int((count * 1.6) ** 0.5)))
        ROWS   = (count + COLS - 1) // COLS
        CELL_W = 270; CELL_H = 130
        PAD    = 50
        W      = max(800, COLS * CELL_W + PAD * 2)
        H      = max(340, ROWS * CELL_H + PAD * 2) + 50

        img = _PILImage.new("RGB", (W * scale, H * scale), (243, 242, 241))
        draw = _PILDraw.Draw(img)
        S = scale

        # Font
        font_sm = font_md = font_bold = None
        try:
            import os
            for p in ["C:/Windows/Fonts/segoeui.ttf",
                      "C:/Windows/Fonts/arial.ttf",
                      "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                      "/System/Library/Fonts/Helvetica.ttc"]:
                if os.path.exists(p):
                    font_sm   = _PILFont.truetype(p, 18 * S // 2)
                    font_md   = _PILFont.truetype(p, 20 * S // 2)
                    font_bold = _PILFont.truetype(p, 22 * S // 2)
                    break
        except Exception:
            pass
        if not font_sm:
            font_sm = font_md = font_bold = _PILFont.load_default()

        pos = {}
        for i, name in enumerate(table_names):
            col = i % COLS; row = i // COLS
            x = PAD + col * CELL_W + (CELL_W - NODE_W) // 2
            y = PAD + row * CELL_H + (CELL_H - NODE_H) // 2
            pos[name] = {"x": x, "y": y, "cx": x + NODE_W//2, "cy": y + NODE_H//2}

        # Draw edges first
        for rel in relationships:
            fp = pos.get(rel.get("fromTable",""))
            tp = pos.get(rel.get("toTable",""))
            if not fp or not tp: continue
            active = rel.get("active", True)
            lc = (0, 120, 212) if active else (161, 159, 157)
            x1, y1 = fp["cx"] * S, fp["cy"] * S
            x2, y2 = tp["cx"] * S, tp["cy"] * S
            # Draw line
            draw.line([(x1, y1), (x2, y2)], fill=lc, width=2 * S // 2)
            # Cardinality label at midpoint
            mx, my = (x1 + x2) // 2, (y1 + y2) // 2
            card = rel.get("cardinality", "")
            if card:
                tw = len(card) * 8 * S // 2
                draw.rectangle([mx - tw - 4, my - 10 * S // 2,
                                 mx + tw + 4, my + 10 * S // 2],
                                fill=(239, 246, 255), outline=lc, width=1)
                draw.text((mx - tw, my - 8 * S // 2), card, font=font_sm, fill=lc)

        # Draw table cards
        col_by_table: dict = {t: [] for t in table_names}
        for rel in relationships:
            for tbl, col in [(rel.get("fromTable",""), rel.get("fromColumn","")),
                             (rel.get("toTable",""),   rel.get("toColumn",""))]:
                if tbl and col and col not in col_by_table.get(tbl, []):
                    col_by_table.setdefault(tbl, []).append(col)

        for name in table_names:
            p = pos[name]
            cols = col_by_table.get(name, [])[:4]
            card_h = NODE_H + len(cols) * 20 + (8 if cols else 0)
            x, y = p["x"] * S, p["y"] * S
            w, h = NODE_W * S, card_h * S

            # Shadow
            draw.rounded_rectangle([x+3, y+3, x+w+3, y+h+3], radius=7*S//2,
                                    fill=(210, 210, 210))
            # Card body
            draw.rounded_rectangle([x, y, x+w, y+h], radius=7*S//2,
                                    fill=(255,255,255), outline=(200,198,196), width=1)
            # Header
            draw.rounded_rectangle([x, y, x+w, y+NODE_H*S], radius=7*S//2,
                                    fill=(16, 124, 16))
            draw.rectangle([x, y+(NODE_H-8)*S, x+w, y+NODE_H*S], fill=(16,124,16))

            # Table name
            short = name if len(name) <= 22 else name[:20] + "…"
            draw.text((x + 40*S//2, y + 14*S//2), short,
                      font=font_bold, fill=(255,255,255))

            # Column rows
            for ci, col in enumerate(cols):
                cy2 = y + (NODE_H + 8 + ci * 20) * S
                draw.line([(x, cy2), (x+w, cy2)], fill=(225,223,221), width=1)
                is_key = "id" in col.lower() or "key" in col.lower()
                tc = (0, 120, 212) if is_key else (59, 58, 57)
                prefix = "# " if is_key else "  "
                draw.text((x + 10*S//2, cy2 + 3*S//2),
                          prefix + col[:25], font=font_sm, fill=tc)

        # Legend
        ly = (H - 24) * S
        draw.line([(PAD*S, ly), ((PAD+26)*S, ly)], fill=(0,120,212), width=2)
        draw.text(((PAD+32)*S, ly - 6*S//2), "Ativo",   font=font_sm, fill=(96,94,92))
        draw.line([((PAD+72)*S, ly), ((PAD+98)*S, ly)], fill=(161,159,157), width=2)
        draw.text(((PAD+104)*S, ly - 6*S//2), "Inativo", font=font_sm, fill=(96,94,92))

        return _pil_png(img)
    except Exception as e:
        return None


def make_arch_png(sources: list, queries: list, graph_edges: list, scale: int = 2) -> bytes | None:
    """Draw the architecture diagram using Pillow. No SVG needed."""
    if not PIL_AVAILABLE or not sources:
        return None
    try:
        CARD_W = 180; CARD_H = 52; GAP_Y = 18
        PBI_W  = 200; PBI_H  = 68; PAD_X = 60; PAD_Y = 40; COL_GAP = 140
        src_count   = len(sources)
        total_src_h = src_count * CARD_H + (src_count - 1) * GAP_Y
        H = max(280, total_src_h + PAD_Y * 2)
        W = PAD_X * 2 + CARD_W + COL_GAP + PBI_W + COL_GAP + CARD_W
        S = scale

        img = _PILImage.new("RGB", (W * S, H * S), (243, 242, 241))
        draw = _PILDraw.Draw(img)

        font_sm = font_bold = None
        try:
            import os
            for p in ["C:/Windows/Fonts/segoeui.ttf",
                      "C:/Windows/Fonts/arial.ttf",
                      "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                      "/System/Library/Fonts/Helvetica.ttc"]:
                if os.path.exists(p):
                    font_sm   = _PILFont.truetype(p, 20 * S // 2)
                    font_bold = _PILFont.truetype(p, 22 * S // 2)
                    break
        except Exception:
            pass
        if not font_sm:
            font_sm = font_bold = _PILFont.load_default()

        SRC_COLORS = {
            "sql": (0,120,212), "excel": (33,115,70), "csv": (33,115,70),
            "postgresql": (51,103,145), "mysql": (0,97,138),
            "json": (242,200,17), "web": (0,120,212), "odata": (0,120,212),
        }
        def src_color(label):
            lbl = label.lower()
            for k, c in SRC_COLORS.items():
                if k in lbl: return c
            return (27,42,56)

        # PBI node position
        pbi_x = (PAD_X + CARD_W + COL_GAP) * S
        pbi_y = ((H - PBI_H) // 2) * S
        pbi_cx = pbi_x + PBI_W * S // 2
        pbi_cy = pbi_y + PBI_H * S // 2

        # Source cards + connection lines
        src_y_start = (H - total_src_h) // 2
        for i, src in enumerate(sources):
            sx = PAD_X * S
            sy = (src_y_start + i * (CARD_H + GAP_Y)) * S
            lbl = src.get("label","Source")
            short = lbl if len(lbl) <= 20 else lbl[:18]+"…"
            r, g, b = src_color(lbl)

            # Card shadow
            draw.rounded_rectangle([sx+3, sy+3, sx+CARD_W*S+3, sy+CARD_H*S+3],
                                    radius=8*S//2, fill=(200,200,200))
            # Card body
            draw.rounded_rectangle([sx, sy, sx+CARD_W*S, sy+CARD_H*S],
                                    radius=8*S//2, fill=(r,g,b))
            # Highlight strip
            draw.rectangle([sx, sy, sx+CARD_W*S, sy+6*S//2],
                            fill=(min(r+40,255), min(g+40,255), min(b+40,255)))
            draw.text((sx + 16*S//2, sy + CARD_H*S//2 - 10*S//2),
                      short, font=font_bold, fill=(255,255,255))

            # Bezier-like connection (straight + elbow)
            src_rx = sx + CARD_W * S
            src_cy2 = sy + CARD_H * S // 2
            mx = (src_rx + pbi_x) // 2
            # Draw polyline: src_right → mid → pbi_left
            draw.line([(src_rx, src_cy2), (mx, src_cy2), (mx, pbi_cy), (pbi_x, pbi_cy)],
                      fill=(0, 120, 212, 160), width=max(1, S//2))

        # PBI Dataset node
        # Gold top/bottom bars
        draw.rounded_rectangle([pbi_x, pbi_y, pbi_x+PBI_W*S, pbi_y+PBI_H*S],
                                radius=10*S//2, fill=(27,42,56))
        draw.rectangle([pbi_x, pbi_y, pbi_x+PBI_W*S, pbi_y+8*S//2],
                        fill=(242,200,17))
        draw.rectangle([pbi_x, pbi_y+(PBI_H-6)*S//2, pbi_x+PBI_W*S, pbi_y+PBI_H*S],
                        fill=(242,200,17))
        draw.text((pbi_x + 16*S//2, pbi_y + 12*S//2),
                  "Power BI Dataset", font=font_bold, fill=(242,200,17))
        draw.text((pbi_x + 40*S//2, pbi_y + 32*S//2),
                  f"{len(sources)} fonte(s)", font=font_sm, fill=(200,200,200))

        return _pil_png(img)
    except Exception:
        return None


def build_erd_svg(relationships: list) -> str:
    """Generate a Power BI–styled ERD SVG from a list of relationship dicts."""
    # Collect unique table names
    table_names = []
    seen_t = set()
    for rel in relationships:
        for t in [rel.get("fromTable", ""), rel.get("toTable", "")]:
            if t and t not in seen_t:
                seen_t.add(t)
                table_names.append(t)
    table_names.sort()

    if not table_names:
        return ""

    count  = len(table_names)
    NODE_W = 190
    NODE_H = 46
    COLS   = min(count, max(2, int((count * 1.6) ** 0.5)))
    ROWS   = (count + COLS - 1) // COLS
    CELL_W = 270
    CELL_H = 120
    PAD_X  = 50
    PAD_Y  = 50
    SVG_W  = max(800, COLS * CELL_W + PAD_X * 2)
    SVG_H  = max(340, ROWS * CELL_H + PAD_Y * 2) + 50

    # Position map
    pos = {}
    for i, name in enumerate(table_names):
        col = i % COLS
        row = i // COLS
        pos[name] = {
            "x":  PAD_X + col * CELL_W + (CELL_W - NODE_W) // 2,
            "y":  PAD_Y + row * CELL_H + (CELL_H - NODE_H) // 2,
            "cx": PAD_X + col * CELL_W + CELL_W // 2,
            "cy": PAD_Y + row * CELL_H + CELL_H // 2,
        }

    # Crow's-foot markers
    defs_parts = []
    edge_parts = []
    label_parts = []

    CARD_MAP = {"1:M": ("one", "many"), "M:1": ("many", "one"),
                "1:1": ("one", "one"), "M:M": ("many", "many")}

    def marker_def(mid, kind, color):
        if kind == "many":
            return (
                f'<marker id="{mid}" viewBox="-2 -6 14 12" refX="11" refY="0"'
                f' markerWidth="14" markerHeight="12" orient="auto">'
                f'<line x1="0" y1="-5" x2="10" y2="0" stroke="{color}" stroke-width="1.8"/>'
                f'<line x1="0" y1="5"  x2="10" y2="0" stroke="{color}" stroke-width="1.8"/>'
                f'<line x1="0" y1="0"  x2="10" y2="0" stroke="{color}" stroke-width="1.8"/>'
                f'</marker>'
            )
        return (
            f'<marker id="{mid}" viewBox="-2 -6 12 12" refX="10" refY="0"'
            f' markerWidth="12" markerHeight="12" orient="auto">'
            f'<line x1="7" y1="-5" x2="7" y2="5" stroke="{color}" stroke-width="1.8"/>'
            f'<line x1="3" y1="-5" x2="3" y2="5" stroke="{color}" stroke-width="1.8"/>'
            f'</marker>'
        )

    for i, rel in enumerate(relationships):
        fp = pos.get(rel.get("fromTable", ""))
        tp = pos.get(rel.get("toTable", ""))
        if not fp or not tp:
            continue

        active = rel.get("active", True)
        color  = "#0078D4" if active else "#A19F9D"
        dash   = 'stroke-dasharray="6,4"' if not active else ""
        card   = rel.get("cardinality", "1:M")
        from_k, to_k = CARD_MAP.get(card, ("one", "many"))

        mid_f = f"mf{i}"
        mid_t = f"mt{i}"
        defs_parts.append(marker_def(mid_f, from_k, color))
        defs_parts.append(marker_def(mid_t, to_k, color))

        dx = tp["cx"] - fp["cx"]
        dy = tp["cy"] - fp["cy"]
        same_row = abs(dy) < CELL_H * 0.5

        if same_row:
            if dx > 0:
                x1, y1 = fp["x"] + NODE_W, fp["y"] + NODE_H // 2
                x2, y2 = tp["x"],           tp["y"] + NODE_H // 2
            else:
                x1, y1 = fp["x"],           fp["y"] + NODE_H // 2
                x2, y2 = tp["x"] + NODE_W,  tp["y"] + NODE_H // 2
            gap = abs(x2 - x1) * 0.45
            cp1x, cp1y = x1 + (gap if dx > 0 else -gap), y1
            cp2x, cp2y = x2 - (gap if dx > 0 else -gap), y2
        else:
            if dy > 0:
                x1, y1 = fp["x"] + NODE_W // 2, fp["y"] + NODE_H
                x2, y2 = tp["x"] + NODE_W // 2, tp["y"]
            else:
                x1, y1 = fp["x"] + NODE_W // 2, fp["y"]
                x2, y2 = tp["x"] + NODE_W // 2, tp["y"] + NODE_H
            gap = abs(y2 - y1) * 0.45
            cp1x, cp1y = x1, y1 + (gap if dy > 0 else -gap)
            cp2x, cp2y = x2, y2 - (gap if dy > 0 else -gap)

        sw = 2 if active else 1.5
        edge_parts.append(
            f'<path d="M {x1} {y1} C {cp1x} {cp1y}, {cp2x} {cp2y}, {x2} {y2}"'
            f' fill="none" stroke="{color}" stroke-width="{sw}" {dash}'
            f' marker-start="url(#{mid_f})" marker-end="url(#{mid_t})"/>'
        )

        # Cardinality pill at midpoint
        mx = (x1 + x2) / 2
        my = (y1 + y2) / 2
        lw = len(card) * 8 + 16
        lh = 20
        cf = rel.get("crossFilter", "")
        cf_icon = "⇄" if "both" in cf.lower() else "→"
        label_parts.append(
            f'<g transform="translate({mx - lw/2:.1f},{my - lh/2 - 2:.1f})">'
            f'<rect width="{lw}" height="{lh}" rx="4"'
            f' fill="{"#EFF6FF" if active else "#F3F2F1"}" stroke="{color}" stroke-width="1" opacity="0.97"/>'
            f'<text x="{lw/2:.1f}" y="{lh/2 + 4.5:.1f}" text-anchor="middle"'
            f' font-family="Consolas,monospace" font-size="10" font-weight="700" fill="{color}">{card}</text>'
            f'</g>'
            f'<text x="{mx:.1f}" y="{my + lh/2 + 11:.1f}" text-anchor="middle"'
            f' font-family="Segoe UI,Arial,sans-serif" font-size="9" fill="{color}" opacity="0.8">{cf_icon}</text>'
        )

    # Table node cards
    col_by_table: dict[str, list[str]] = {t: [] for t in table_names}
    for rel in relationships:
        ft, fc = rel.get("fromTable", ""), rel.get("fromColumn", "")
        tt, tc = rel.get("toTable", ""), rel.get("toColumn", "")
        if ft and fc and fc not in col_by_table.get(ft, []):
            col_by_table.setdefault(ft, []).append(fc)
        if tt and tc and tc not in col_by_table.get(tt, []):
            col_by_table.setdefault(tt, []).append(tc)

    card_parts = []
    for name in table_names:
        p = pos[name]
        cols = col_by_table.get(name, [])[:4]
        card_h = NODE_H + max(0, len(cols)) * 19 + (8 if cols else 0)
        short = name if len(name) <= 22 else name[:20] + "…"

        col_rows_svg = ""
        for ci, col in enumerate(cols):
            cy = NODE_H + 6 + ci * 19
            is_key = "id" in col.lower() or "key" in col.lower()
            key_prefix = "🔑 " if is_key else ""
            col_color = "#0078D4" if is_key else "#3B3A39"
            col_short = col if len(col) <= 25 else col[:23] + "…"
            col_rows_svg += (
                f'<line x1="0" y1="{cy-1}" x2="{NODE_W}" y2="{cy-1}" stroke="#E1DFDD" stroke-width="0.8"/>'
                f'<text x="10" y="{cy+12}" font-family="Segoe UI,Arial,sans-serif" font-size="10" fill="{col_color}">'
                f'{key_prefix}{col_short}</text>'
            )

        card_parts.append(
            f'<g transform="translate({p["x"]},{p["y"]})">'
            # shadow
            f'<rect x="3" y="3" width="{NODE_W}" height="{card_h}" rx="7" fill="rgba(0,0,0,0.10)"/>'
            # white body
            f'<rect width="{NODE_W}" height="{card_h}" rx="7" fill="white" stroke="#C8C6C4" stroke-width="0.8"/>'
            # header fill
            f'<rect width="{NODE_W}" height="{NODE_H}" rx="7" fill="#107C10"/>'
            f'<rect y="{NODE_H-7}" width="{NODE_W}" height="7" fill="#107C10"/>'
            # icon box
            f'<rect x="8" y="9" width="26" height="26" rx="5" fill="rgba(255,255,255,0.18)"/>'
            f'<text x="21" y="27" text-anchor="middle" font-family="Segoe UI,Arial,sans-serif" font-size="13">🗄</text>'
            # table name
            f'<text x="{NODE_W//2 + 8}" y="29" text-anchor="middle"'
            f' font-family="Segoe UI,Arial,sans-serif" font-size="11" font-weight="700" fill="white">{short}</text>'
            f'{col_rows_svg}'
            f'</g>'
        )

    # Legend
    legend = (
        f'<g transform="translate({PAD_X},{SVG_H - 30})">'
        f'<line x1="0" y1="6" x2="26" y2="6" stroke="#0078D4" stroke-width="2"/>'
        f'<text x="32" y="10" font-family="Segoe UI,Arial,sans-serif" font-size="10" fill="#605E5C">Ativo</text>'
        f'<line x1="72" y1="6" x2="98" y2="6" stroke="#A19F9D" stroke-width="1.5" stroke-dasharray="5,3"/>'
        f'<text x="104" y="10" font-family="Segoe UI,Arial,sans-serif" font-size="10" fill="#605E5C">Inativo</text>'
        f'</g>'
    )

    return (
        f'<svg viewBox="0 0 {SVG_W} {SVG_H}" width="{SVG_W}" height="{SVG_H}"'
        f' xmlns="http://www.w3.org/2000/svg"'
        f' style="background:#F3F2F1;font-family:Segoe UI,Arial,sans-serif">'
        f'<rect width="{SVG_W}" height="{SVG_H}" fill="#F3F2F1"/>'
        f'<defs>{"".join(defs_parts)}</defs>'
        f'{"".join(edge_parts)}'
        f'{"".join(card_parts)}'
        f'{"".join(label_parts)}'
        f'{legend}'
        f'</svg>'
    )


def build_architecture_svg(sources: list, queries: list, edges: list) -> str:
    """Generate an architecture diagram SVG: Sources → [Power Query] → Power BI Dataset."""
    source_nodes = [n for n in sources]
    query_nodes  = [n for n in queries]

    if not source_nodes:
        return ""

    # Map source id → queries
    src_to_q: dict[str, list] = {s["id"]: [] for s in source_nodes}
    for e in edges:
        if e.get("from") in src_to_q:
            q = next((n for n in query_nodes if n["id"] == e.get("to")), None)
            if q:
                src_to_q[e["from"]].append(q)

    CARD_W   = 160
    CARD_H   = 52
    PBI_W    = 180
    PBI_H    = 64
    GAP_Y    = 18
    PAD_X    = 50
    PAD_Y    = 40
    COL_GAP  = 120

    src_count = len(source_nodes)
    total_src_h = src_count * CARD_H + (src_count - 1) * GAP_Y
    SVG_H = max(260, total_src_h + PAD_Y * 2)
    SVG_W = PAD_X * 2 + CARD_W + COL_GAP + PBI_W + COL_GAP + CARD_W

    parts  = []
    lines  = []

    # PBI Dataset node (center-right)
    pbi_x = PAD_X + CARD_W + COL_GAP
    pbi_y = (SVG_H - PBI_H) // 2

    # Source cards (left column)
    src_y_start = (SVG_H - total_src_h) // 2
    SRC_COLORS = {
        "sql": "#0078D4", "excel": "#217346", "csv": "#217346",
        "postgresql": "#336791", "mysql": "#00618A", "sharepoint": "#0078D4",
        "web": "#0078D4", "json": "#F2C811", "odata": "#0078D4",
    }

    def src_color(label):
        lbl = label.lower()
        for k, c in SRC_COLORS.items():
            if k in lbl:
                return c
        return "#1B2A38"

    for i, src in enumerate(source_nodes):
        sx = PAD_X
        sy = src_y_start + i * (CARD_H + GAP_Y)
        lbl = src.get("label", "Source")
        short = lbl if len(lbl) <= 20 else lbl[:18] + "…"
        color = src_color(lbl)

        # card
        parts.append(
            f'<g transform="translate({sx},{sy})">'
            f'<rect x="2" y="2" width="{CARD_W}" height="{CARD_H}" rx="8" fill="rgba(0,0,0,0.10)"/>'
            f'<rect width="{CARD_W}" height="{CARD_H}" rx="8" fill="{color}"/>'
            f'<rect width="{CARD_W}" height="6" rx="0" fill="rgba(255,255,255,0.15)" y="0"/>'
            f'<rect width="4" height="{CARD_H}" rx="0" fill="rgba(255,255,255,0.25)" x="0"/>'
            f'<text x="{CARD_W//2}" y="{CARD_H//2 + 5}" text-anchor="middle"'
            f' font-family="Segoe UI,Arial,sans-serif" font-size="11" font-weight="600" fill="white">{short}</text>'
            f'</g>'
        )

        # connection line: src card right edge → PBI node left edge
        src_cx = sx + CARD_W
        src_cy = sy + CARD_H // 2
        pbi_lx = pbi_x
        pbi_cy = pbi_y + PBI_H // 2

        # bezier
        mid_x = (src_cx + pbi_lx) // 2
        lines.append(
            f'<path d="M {src_cx} {src_cy} C {mid_x} {src_cy}, {mid_x} {pbi_cy}, {pbi_lx} {pbi_cy}"'
            f' fill="none" stroke="#0078D4" stroke-width="1.8" opacity="0.55"'
            f' stroke-dasharray="6,3"/>'
        )

    # PBI Dataset card
    parts.append(
        f'<g transform="translate({pbi_x},{pbi_y})">'
        f'<rect x="3" y="3" width="{PBI_W}" height="{PBI_H}" rx="10" fill="rgba(0,0,0,0.13)"/>'
        f'<rect width="{PBI_W}" height="{PBI_H}" rx="10" fill="#1B2A38"/>'
        f'<rect width="{PBI_W}" height="8" rx="0" fill="#F2C811" y="0"/>'
        f'<rect width="{PBI_W}" height="8" rx="0" fill="#F2C811" y="{PBI_H-8}"/>'
        f'<text x="{PBI_W//2}" y="{PBI_H//2 - 4}" text-anchor="middle"'
        f' font-family="Segoe UI,Arial,sans-serif" font-size="10" font-weight="700" fill="#F2C811">⚡ Power BI</text>'
        f'<text x="{PBI_W//2}" y="{PBI_H//2 + 12}" text-anchor="middle"'
        f' font-family="Segoe UI,Arial,sans-serif" font-size="11" font-weight="600" fill="white">Dataset</text>'
        f'</g>'
    )

    return (
        f'<svg viewBox="0 0 {SVG_W} {SVG_H}" width="{SVG_W}" height="{SVG_H}"'
        f' xmlns="http://www.w3.org/2000/svg"'
        f' style="background:#F3F2F1">'
        f'<rect width="{SVG_W}" height="{SVG_H}" fill="#F3F2F1"/>'
        f'{"".join(lines)}'
        f'{"".join(parts)}'
        f'</svg>'
    )


def insert_svg_image(doc, svg_string: str, width_inches: float = 6.5, caption: str = "") -> bool:
    """Render SVG to PNG and insert into the document. Returns True on success."""
    if not svg_string:
        return False
    png_bytes = svg_to_png_bytes(svg_string, scale=2.0)
    if not png_bytes:
        return False
    buf = BytesIO(png_bytes)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run().add_picture(buf, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cap_run = cap.add_run(caption)
        cap_run.font.size = Pt(8.5)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(96, 94, 92)
    return True


DOC_TEXT = {
    "pt-BR": {
        "metrics": {
            "sources": "Fontes",
            "queries": "Queries",
            "tables": "Tabelas",
            "measures": "Medidas",
            "calc_columns": "Colunas calculadas",
            "relationships": "Relacionamentos",
            "pages": "Páginas",
            "visuals": "Visuais",
        },
        "sections": [
            "Visão geral",
            "Fontes de dados",
            "Power Query",
            "Modelo semântico",
            "Dicionário de dados",
            "Medidas DAX",
            "Colunas calculadas",
            "Relacionamentos",
            "Páginas e visuais",
            "Linhagem técnica",
            "Diagnósticos",
        ],
        "cover_subtitle": "Documentação técnica do painel",
        "generated_at": "Gerado em",
        "toc": "Sumário",
        "overview_banner": "Visão geral do relatório",
        "overview_text": "Documentação gerada automaticamente pelo BI Flow Mapper a partir dos metadados locais do arquivo PBIX.",
        "file": "Arquivo",
        "analysis_source": "Origem da análise",
        "total_nodes": "Total de nós no mapa",
        "total_edges": "Total de arestas de linhagem",
        "data_sources_banner": "Conectores e origens de dados detectados",
        "source_headers": ["Fonte", "Padrão detectado", "Caminho / servidor"],
        "no_sources": "Nenhuma fonte detectada.",
        "arch_caption": "Diagrama de arquitetura — fontes conectadas ao dataset Power BI",
        "sources_detected": "fonte(s) detectada(s):",
        "connection_details": "2.1 Detalhes de conexão (TMSCHEMA_DATASOURCES)",
        "connection_headers": ["Tipo", "Nome", "String de conexão"],
        "power_query_banner": "Transformações e consultas M",
        "query_headers": ["Query", "Conexão / caminho", "Resumo técnico"],
        "no_queries": "Nenhuma query encontrada.",
        "semantic_model_banner": "Tabelas carregadas no modelo",
        "table_headers": ["Tabela", "Colunas", "Oculta", "Descrição"],
        "no_tables": "Nenhuma tabela encontrada.",
        "dictionary_banner": "Colunas, tipos e metadados de cada tabela",
        "column_headers": ["Tabela", "Coluna", "Tipo", "Formato", "Oculta", "Expressão"],
        "no_columns": "Nenhuma coluna encontrada nos metadados.",
        "measures_banner": "Medidas calculadas com expressões DAX",
        "measure_headers": ["Tabela", "Medida"],
        "no_measures": "Nenhuma medida encontrada.",
        "calc_columns_banner": "Colunas derivadas por expressões DAX",
        "calc_headers": ["Tabela", "Coluna"],
        "no_calc_columns": "Nenhuma coluna calculada encontrada.",
        "relationships_banner": "Vínculos entre tabelas do modelo semântico",
        "relationship_headers": ["Tabela origem", "Coluna origem", "Tabela destino", "Coluna destino", "Card.", "Filtro cruzado", "Ativo"],
        "no_relationships": "Nenhum relacionamento encontrado.",
        "erd_caption": "Diagrama entidade-relacionamento — modelo semântico Power BI",
        "pages_banner": "Páginas do relatório e seus elementos visuais",
        "page_headers": ["#", "Página", "Visuais", "Canvas (px)"],
        "no_pages": "Nenhuma página encontrada.",
        "visual_headers": ["Visual", "Tipo", "Campos / referências"],
        "no_visuals": "Nenhum visual encontrado.",
        "lineage_banner": "Grafo completo de dependências do pipeline",
        "lineage_headers": ["Origem", "Relacao", "Destino"],
        "no_lineage": "Nenhuma aresta de linhagem encontrada.",
        "diagnostics_banner": "Alertas e informações técnicas da análise",
        "relationship_columns": "Colunas disponíveis na tabela de relacionamentos:",
        "rel_from": "Origem",
        "rel_to": "Destino",
        "cardinality": "Cardinalidade",
        "cross_filter": "Filtro cruzado",
        "relationship_example": "Exemplo de relacionamento detectado — ",
        "message": "Mensagem",
        "no_diagnostics": "Nenhum diagnóstico registrado.",
        "yes": "Sim",
        "no": "Não",
        "dash": "—",
        "m_code_heading": "3.1 Código M por consulta",
        "summary": "Resumo",
        "technical_expressions": "Expressões técnicas",
        "additional_records_omitted": "{count} registros adicionais foram omitidos nesta seção para manter o documento legível.",
        "no_records": "Nenhum registro encontrado.",
        "record": "Registro",
        "truncated": "truncado",
        "connector": "Conector",
        "steps": "Etapas",
        "functions": "Funções",
        "datatype_map": {
            "2": "texto", "3": "decimal", "4": "inteiro", "5": "decimal",
            "6": "moeda", "7": "data", "8": "booleano", "9": "binário",
            "10": "variant", "17": "inteiro",
            "string": "texto", "int64": "inteiro", "double": "decimal",
            "boolean": "booleano", "datetime": "data/hora", "binary": "binário",
        },
    },
    "en-US": {
        "metrics": {
            "sources": "Sources",
            "queries": "Queries",
            "tables": "Tables",
            "measures": "Measures",
            "calc_columns": "Calculated columns",
            "relationships": "Relationships",
            "pages": "Pages",
            "visuals": "Visuals",
        },
        "sections": [
            "Overview",
            "Data sources",
            "Power Query",
            "Semantic model",
            "Data dictionary",
            "DAX measures",
            "Calculated columns",
            "Relationships",
            "Pages and visuals",
            "Technical lineage",
            "Diagnostics",
        ],
        "cover_subtitle": "Technical report documentation",
        "generated_at": "Generated at",
        "toc": "Table of contents",
        "overview_banner": "Report overview",
        "overview_text": "Documentation automatically generated by BI Flow Mapper from the local metadata in the PBIX file.",
        "file": "File",
        "analysis_source": "Analysis source",
        "total_nodes": "Total nodes in the map",
        "total_edges": "Total lineage edges",
        "data_sources_banner": "Detected connectors and data sources",
        "source_headers": ["Source", "Detected pattern", "Path / server"],
        "no_sources": "No sources detected.",
        "arch_caption": "Architecture diagram — sources connected to the Power BI dataset",
        "sources_detected": "source(s) detected:",
        "connection_details": "2.1 Connection details (TMSCHEMA_DATASOURCES)",
        "connection_headers": ["Type", "Name", "Connection string"],
        "power_query_banner": "M transformations and queries",
        "query_headers": ["Query", "Connection / path", "Technical summary"],
        "no_queries": "No queries found.",
        "semantic_model_banner": "Tables loaded into the model",
        "table_headers": ["Table", "Columns", "Hidden", "Description"],
        "no_tables": "No tables found.",
        "dictionary_banner": "Columns, types, and metadata for each table",
        "column_headers": ["Table", "Column", "Type", "Format", "Hidden", "Expression"],
        "no_columns": "No columns found in the metadata.",
        "measures_banner": "Calculated measures with DAX expressions",
        "measure_headers": ["Table", "Measure"],
        "no_measures": "No measures found.",
        "calc_columns_banner": "Columns derived from DAX expressions",
        "calc_headers": ["Table", "Column"],
        "no_calc_columns": "No calculated columns found.",
        "relationships_banner": "Links between semantic model tables",
        "relationship_headers": ["Source table", "Source column", "Target table", "Target column", "Card.", "Cross filter", "Active"],
        "no_relationships": "No relationships found.",
        "erd_caption": "Entity relationship diagram — Power BI semantic model",
        "pages_banner": "Report pages and their visual elements",
        "page_headers": ["#", "Page", "Visuals", "Canvas (px)"],
        "no_pages": "No pages found.",
        "visual_headers": ["Visual", "Type", "Fields / references"],
        "no_visuals": "No visuals found.",
        "lineage_banner": "Complete dependency graph for the pipeline",
        "lineage_headers": ["Source", "Relationship", "Target"],
        "no_lineage": "No lineage edges found.",
        "diagnostics_banner": "Warnings and technical information from the analysis",
        "relationship_columns": "Available columns in the relationships table:",
        "rel_from": "From",
        "rel_to": "To",
        "cardinality": "Cardinality",
        "cross_filter": "Cross filter",
        "relationship_example": "Detected relationship example — ",
        "message": "Message",
        "no_diagnostics": "No diagnostics recorded.",
        "yes": "Yes",
        "no": "No",
        "dash": "—",
        "m_code_heading": "3.1 M code by query",
        "summary": "Summary",
        "technical_expressions": "Technical expressions",
        "additional_records_omitted": "{count} additional records were omitted in this section to keep the document readable.",
        "no_records": "No records found.",
        "record": "Record",
        "truncated": "truncated",
        "connector": "Connector",
        "steps": "Steps",
        "functions": "Functions",
        "datatype_map": {
            "2": "text", "3": "decimal", "4": "integer", "5": "decimal",
            "6": "currency", "7": "date", "8": "boolean", "9": "binary",
            "10": "variant", "17": "integer",
            "string": "text", "int64": "integer", "double": "decimal",
            "boolean": "boolean", "datetime": "date/time", "binary": "binary",
        },
    },
}


def normalize_doc_locale(locale: str = "") -> str:
    normalized = str(locale or "").strip().lower()
    if normalized.startswith("en"):
        return "en-US"
    return "pt-BR"


def doc_text(locale: str = ""):
    return DOC_TEXT[normalize_doc_locale(locale)]


def build_documentation_docx(path: Path, file_name: str = "", locale: str = "pt-BR") -> bytes:
    if not load_pbixray():
        raise RuntimeError(f"pbixray nao esta instalado ou nao carregou: {PBIXRAY_IMPORT_ERROR}")
    if not load_docx():
        raise RuntimeError(f"python-docx nao esta instalado ou nao carregou: {DOCX_IMPORT_ERROR}")
    load_pillow()

    doc_labels = doc_text(locale)
    metric_labels = doc_labels["metrics"]
    diagnostics = []
    model = PBIXRay(str(path))
    graph = analyze_pbix(path)

    power_query = records_from(model, "power_query", diagnostics)
    datasource_records = records_from(model, "tmschema_datasources", diagnostics)
    measures = records_from(model, "dax_measures", diagnostics)
    calc_columns = records_from(model, "dax_columns", diagnostics)
    schema = records_from(model, "schema", diagnostics)
    semantic_tables = records_from(model, "tmschema_tables", diagnostics)
    tmschema_columns = records_from(model, "tmschema_columns", diagnostics)

    nodes = graph.get("nodes", [])
    edges = graph.get("edges", [])
    relationships = graph.get("relationships", [])
    pages = graph.get("pages", [])
    warnings = list(graph.get("warnings", [])) + diagnostics

    sources = [node for node in nodes if node.get("type") == "source"]
    queries = [node for node in nodes if node.get("type") == "query"]
    tables = [node for node in nodes if node.get("type") == "model"]
    measure_nodes = [node for node in nodes if node.get("type") == "measure"]
    calc_column_nodes = [node for node in nodes if node.get("type") == "calc_column"]
    visuals = [node for node in nodes if node.get("type") == "visual"]

    doc = Document()
    configure_document_styles(doc)
    add_cover(doc, file_name or path.name, {
        metric_labels["sources"]: len(sources),
        metric_labels["queries"]: len(queries),
        metric_labels["tables"]: len(tables),
        metric_labels["measures"]: len(measure_nodes),
        metric_labels["calc_columns"]: len(calc_column_nodes),
        metric_labels["relationships"]: len(relationships),
        metric_labels["pages"]: len(pages),
        metric_labels["visuals"]: len(visuals),
    }, doc_labels)
    add_table_of_contents(doc, doc_labels["sections"], doc_labels)

    add_doc_heading(doc, f"1. {doc_labels['sections'][0]}", level=1)
    add_section_banner(doc, doc_labels["overview_banner"], color="#0078D4", icon="📋")
    add_paragraph(
        doc,
        doc_labels["overview_text"],
    )
    add_key_value_table(doc, [
        (doc_labels["file"], file_name or path.name),
        (doc_labels["generated_at"], datetime.now().strftime("%d/%m/%Y %H:%M:%S")),
        (doc_labels["analysis_source"], "PBIXRay"),
        (doc_labels["total_nodes"], len(nodes)),
        (doc_labels["total_edges"], len(edges)),
    ])

    add_doc_heading(doc, f"2. {doc_labels['sections'][1]}", level=1)
    add_section_banner(doc, doc_labels["data_sources_banner"], color="#0078D4", icon="🔌")
    source_rows = []
    for source in sources:
        meta = source.get("meta", {})
        # Connection path from architecture view: extract server/db/path
        datasource_meta = meta.get("datasource") or {}
        connection_string = ""
        if isinstance(datasource_meta, dict):
            connection_string = (
                datasource_meta.get("ConnectionString", "")
                or datasource_meta.get("Location", "")
                or datasource_meta.get("Account", "")
                or datasource_meta.get("Path", "")
                or ""
            )
        if not connection_string:
            connection_string = meta.get("connectionPath") or meta.get("doc") or ""
        source_rows.append([
            source.get("label", ""),
            meta.get("pattern", ""),
            str(connection_string) if connection_string else doc_labels["dash"],
        ])
    add_records_table(doc, doc_labels["source_headers"], source_rows, empty=doc_labels["no_sources"], labels=doc_labels)

    # ── Diagrama de arquitetura ──────────────────────────────────────────────
    arch_png = make_arch_png(sources, queries, edges)
    if arch_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(BytesIO(arch_png), width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(doc_labels["arch_caption"])
        cr.font.size = Pt(8.5); cr.font.italic = True; cr.font.color.rgb = RGBColor(96,94,92)
    elif sources:
        add_paragraph(doc, f"{len(sources)} {doc_labels['sources_detected']} " + ", ".join(s.get('label','') for s in sources))

    if datasource_records:
        add_doc_heading(doc, doc_labels["connection_details"], level=2)
        arch_rows = []
        for ds in datasource_records[:30]:
            conn_str = (
                ds.get("ConnectionString", "")
                or ds.get("Location", "")
                or ds.get("Account", "")
                or ds.get("Path", "")
                or ""
            )
            kind = ds.get("Kind", ds.get("Type", ds.get("SourceType", "")))
            name = ds.get("Name", ds.get("ContentPath", ""))
            arch_rows.append([
                str(kind) if kind else doc_labels["dash"],
                str(name) if name else doc_labels["dash"],
                str(conn_str) if conn_str else doc_labels["dash"],
            ])
        if arch_rows:
            add_records_table(doc, doc_labels["connection_headers"], arch_rows, empty="", labels=doc_labels)
        else:
            add_generic_records_table(doc, datasource_records, max_rows=30, labels=doc_labels)

    add_doc_heading(doc, f"3. {doc_labels['sections'][2]}", level=1)
    add_section_banner(doc, doc_labels["power_query_banner"], color="#F2C811", icon="⚙️")
    query_rows = []
    for query in queries:
        meta = query.get("meta", {})
        query_rows.append([
            query.get("label", ""),
            meta.get("connectionPath", ""),
            summarize_m_expression(meta.get("expression", ""), doc_labels),
        ])
    add_records_table(doc, doc_labels["query_headers"], query_rows, empty=doc_labels["no_queries"], labels=doc_labels)

    add_power_query_evidence(doc, power_query, doc_labels)

    add_doc_heading(doc, f"4. {doc_labels['sections'][3]}", level=1)
    add_section_banner(doc, doc_labels["semantic_model_banner"], color="#107C10", icon="🗄️")
    table_rows = build_table_documentation_rows(tables, semantic_tables, schema, doc_labels)
    add_records_table(doc, doc_labels["table_headers"], table_rows, empty=doc_labels["no_tables"], labels=doc_labels)

    add_doc_heading(doc, f"5. {doc_labels['sections'][4]}", level=1)
    add_section_banner(doc, doc_labels["dictionary_banner"], color="#107C10", icon="📖")
    column_rows = build_column_rows(schema, tmschema_columns, semantic_tables, doc_labels)
    add_records_table(
        doc,
        doc_labels["column_headers"],
        column_rows,
        empty=doc_labels["no_columns"],
        max_rows=250,
        labels=doc_labels,
    )

    add_doc_heading(doc, f"6. {doc_labels['sections'][5]}", level=1)
    add_section_banner(doc, doc_labels["measures_banner"], color="#D83B01", icon="📐")
    measure_rows = build_expression_rows(measures, measure_nodes, kind="measure")
    add_expression_inventory(doc, doc_labels["measure_headers"], measure_rows, empty=doc_labels["no_measures"], labels=doc_labels)

    add_doc_heading(doc, f"7. {doc_labels['sections'][6]}", level=1)
    add_section_banner(doc, doc_labels["calc_columns_banner"], color="#9B5094", icon="🔢")
    calc_rows = build_expression_rows(calc_columns, calc_column_nodes, kind="calc_column")
    add_expression_inventory(doc, doc_labels["calc_headers"], calc_rows, empty=doc_labels["no_calc_columns"], labels=doc_labels)

    add_doc_heading(doc, f"8. {doc_labels['sections'][7]}", level=1)
    add_section_banner(doc, doc_labels["relationships_banner"], color="#0078D4", icon="🔗")
    relationship_rows = [[
        rel.get("fromTable", ""),
        rel.get("fromColumn", ""),
        rel.get("toTable", ""),
        rel.get("toColumn", ""),
        rel.get("cardinality", ""),
        rel.get("crossFilter", ""),
        doc_labels["yes"] if rel.get("active", True) else doc_labels["no"],
    ] for rel in relationships]
    add_records_table(
        doc,
        doc_labels["relationship_headers"],
        relationship_rows,
        empty=doc_labels["no_relationships"],
        max_rows=200,
        labels=doc_labels,
    )

    # ── Diagrama ERD ─────────────────────────────────────────────────────────
    erd_png = make_erd_png(relationships)
    if erd_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(BytesIO(erd_png), width=Inches(6.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(doc_labels["erd_caption"])
        cr.font.size = Pt(8.5); cr.font.italic = True; cr.font.color.rgb = RGBColor(96,94,92)

    add_doc_heading(doc, f"9. {doc_labels['sections'][8]}", level=1)
    add_section_banner(doc, doc_labels["pages_banner"], color="#8764B8", icon="📄")
    page_rows = [[page.get("ordinal", 0) + 1, page.get("name", ""), page.get("visualCount", 0), f"{page.get('width', '')} x {page.get('height', '')}"] for page in pages]
    add_records_table(doc, doc_labels["page_headers"], page_rows, empty=doc_labels["no_pages"], labels=doc_labels)

    visual_rows = []
    for visual in visuals:
        meta = visual.get("meta", {})
        refs_raw = meta.get("refs", []) or []
        # Clean and deduplicate refs: remove generic/noisy values, title-case
        seen_refs = set()
        clean_refs = []
        for ref in refs_raw:
            ref_str = str(ref).strip()
            # Skip very short, numeric-only or UUID-like values
            if len(ref_str) < 2:
                continue
            ref_lower = ref_str.lower()
            if ref_lower in seen_refs:
                continue
            seen_refs.add(ref_lower)
            clean_refs.append(ref_str)
        refs_display = "\n".join(clean_refs) if clean_refs else doc_labels["dash"]
        visual_rows.append([
            visual.get("label", ""),
            meta.get("visualType", ""),
            refs_display,
        ])
    add_records_table(doc, doc_labels["visual_headers"], visual_rows, empty=doc_labels["no_visuals"], max_rows=120, labels=doc_labels)

    add_doc_heading(doc, f"10. {doc_labels['sections'][9]}", level=1)
    add_section_banner(doc, doc_labels["lineage_banner"], color="#1B2A38", icon="🔀")
    edge_rows = []
    node_labels = {node.get("id"): node.get("label", node.get("id", "")) for node in nodes}
    for item in edges:
        edge_rows.append([
            node_labels.get(item.get("from"), item.get("from", "")),
            item.get("label", ""),
            node_labels.get(item.get("to"), item.get("to", "")),
        ])
    add_records_table(doc, doc_labels["lineage_headers"], edge_rows, empty=doc_labels["no_lineage"], max_rows=300, labels=doc_labels)

    add_doc_heading(doc, f"11. {doc_labels['sections'][10]}", level=1)
    add_section_banner(doc, doc_labels["diagnostics_banner"], color="#605E5C", icon="🔍")
    readable_warnings = []
    for w in unique_texts(warnings):
        if not w:
            continue
        if w.startswith("relationships columns:"):
            cols_match = re.search(r"\[(.+)\]", w)
            if cols_match:
                cols = [c.strip().strip("'") for c in cols_match.group(1).split(",")]
                readable_warnings.append(f"{doc_labels['relationship_columns']} {', '.join(cols)}")
            else:
                readable_warnings.append(w)
        elif w.startswith("relationships first row:"):
            try:
                dict_str = w[len("relationships first row: "):]
                row_dict = {}
                for m in re.finditer(r"'([^']+)':\s*(?:'([^']*)'|(\d+)|None)", dict_str):
                    key, sv, nv = m.group(1), m.group(2), m.group(3)
                    row_dict[key] = sv if sv is not None else (nv if nv is not None else "")
                parts = []
                if row_dict.get("FromTableName"):
                    parts.append(f"{doc_labels['rel_from']}: {row_dict['FromTableName']}[{row_dict.get('FromColumnName', '')}]")
                if row_dict.get("ToTableName"):
                    parts.append(f"{doc_labels['rel_to']}: {row_dict['ToTableName']}[{row_dict.get('ToColumnName', '')}]")
                if row_dict.get("Cardinality"):
                    parts.append(f"{doc_labels['cardinality']}: {row_dict['Cardinality']}")
                if row_dict.get("CrossFilteringBehavior"):
                    parts.append(f"{doc_labels['cross_filter']}: {row_dict['CrossFilteringBehavior']}")
                if parts:
                    readable_warnings.append(doc_labels["relationship_example"] + " | ".join(parts))
                else:
                    readable_warnings.append(w)
            except Exception:
                readable_warnings.append(w)
        else:
            readable_warnings.append(w)
    warning_rows = [[w] for w in readable_warnings if w]
    add_records_table(doc, [doc_labels["message"]], warning_rows, empty=doc_labels["no_diagnostics"], max_rows=120, labels=doc_labels)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def configure_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("BI Flow Mapper")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(119, 134, 150)

    for style_name, size, color in [
        ("Title", 24, "0B1F33"),
        ("Heading 1", 15, "0B1F33"),
        ("Heading 2", 11, "1E6D85"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc, file_name, metrics, labels=None):
    labels = labels or doc_text("pt-BR")
    # ── Gold top stripe (Power BI identity) ─────────────────────────────────
    def _make_stripe() -> bytes | None:
        if not PIL_AVAILABLE:
            return None
        try:
            img = _PILImage.new("RGB", (1440, 36), (27, 42, 56))
            draw = _PILDraw.Draw(img)
            draw.rectangle([0, 24, 1440, 36], fill=(242, 200, 17))
            return _pil_png(img)
        except Exception:
            return None
    stripe_png = _make_stripe()
    if stripe_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(BytesIO(stripe_png), width=Inches(7.0))

    logo = ROOT / "image" / "icon.png"
    if logo.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.paragraph_format.space_before = Pt(6)
        logo_paragraph.paragraph_format.space_after = Pt(0)
        logo_paragraph.add_run().add_picture(str(logo), width=Inches(0.65))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("BI Flow Mapper")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(11, 31, 51)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle.add_run(labels["cover_subtitle"])
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(30, 109, 133)

    file_paragraph = doc.add_paragraph()
    file_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_paragraph.paragraph_format.space_after = Pt(22)
    file_run = file_paragraph.add_run(file_name)
    file_run.bold = True
    file_run.font.size = Pt(15)
    file_run.font.color.rgb = RGBColor(35, 48, 64)

    add_metric_grid(doc, metrics)
    add_paragraph(doc, f"{labels['generated_at']} {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    # ── Gold bottom stripe ───────────────────────────────────────────────────
    if stripe_png:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(10)
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run().add_picture(BytesIO(stripe_png), width=Inches(7.0))

    doc.add_page_break()


def add_table_of_contents(doc, sections, labels=None):
    labels = labels or doc_text("pt-BR")
    add_doc_heading(doc, labels["toc"], level=1)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    set_table_width(table, [0.55, 6.2])
    for index, title in enumerate(sections, 1):
        cells = table.add_row().cells
        cells[0].text = f"{index:02d}"
        cells[1].text = title
        for cell in cells:
            set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
            if index % 2 == 0:
                shade_cell(cell, "F8FAFC")
        set_cell_font(cells[0], color="1E6D85", bold=True)
        set_cell_font(cells[1], color="1F2937", bold=True)
    doc.add_page_break()


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(str(text or ""))
    return paragraph


def add_section_banner(doc, text: str, color: str = "#0078D4", icon: str = ""):
    """Insert a colored banner below a heading for visual identity."""
    png = make_banner_png(text, color=color, icon=icon)
    if not png:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(BytesIO(png), width=Inches(6.8))


def add_doc_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    run = paragraph.add_run(str(text))
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(15 if level == 1 else 11)
    run.font.color.rgb = RGBColor(11, 31, 51) if level == 1 else RGBColor(30, 109, 133)
    if level == 1:
        add_paragraph_bottom_border(paragraph, "D8DEE8", "8")
    return paragraph


def add_metric_grid(doc, metrics):
    items = list(metrics.items())
    cols = 4
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    for row_index in range(rows):
        for col_index in range(cols):
            cell = table.rows[row_index].cells[col_index]
            set_cell_margins(cell, top=140, start=140, bottom=140, end=140)
            shade_cell(cell, "F4F7FA")
            item_index = row_index * cols + col_index
            if item_index >= len(items):
                cell.text = ""
                continue
            label, value = items[item_index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            value_run = paragraph.add_run(str(value))
            value_run.bold = True
            value_run.font.size = Pt(18)
            value_run.font.color.rgb = RGBColor(11, 31, 51)
            paragraph.add_run("\n")
            label_run = paragraph.add_run(str(label))
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(88, 103, 118)
    doc.add_paragraph("")


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table, color="D8DEE8")
    set_table_width(table, [2.1, 4.9])
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
        shade_cell(cells[0], "F4F7FA")
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        bold_cell(cells[0])
    doc.add_paragraph("")


def add_records_table(doc, headers, rows, empty="", max_rows=120, labels=None):
    labels = labels or doc_text("pt-BR")
    if not rows:
        add_paragraph(doc, empty)
        return

    limited_rows = rows[:max_rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_table_borders(table, color="D8DEE8")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade_cell(cell, "0B1F33")
        set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
        set_cell_font(cell, color="FFFFFF", bold=True)

    for row_index, row in enumerate(limited_rows):
        cells = table.add_row().cells
        for index, value in enumerate(row[:len(headers)]):
            cells[index].text = clip(value, 700, labels)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[index], top=70, start=90, bottom=70, end=90)
            if row_index % 2:
                shade_cell(cells[index], "F8FAFC")
            set_cell_font(cells[index], color="1F2937")

    if len(rows) > max_rows:
        add_paragraph(doc, labels["additional_records_omitted"].format(count=len(rows) - max_rows))
    doc.add_paragraph("")


def add_generic_records_table(doc, records, max_rows=60, labels=None):
    labels = labels or doc_text("pt-BR")
    keys = []
    for record in records:
        for key in record.keys():
            if key not in keys:
                keys.append(key)
        if len(keys) >= 6:
            break
    rows = [[doc_value(record.get(key, "")) for key in keys] for record in records]
    add_records_table(doc, keys or [labels["record"]], rows, empty=labels["no_records"], max_rows=max_rows, labels=labels)


def add_power_query_evidence(doc, power_query, labels=None):
    labels = labels or doc_text("pt-BR")
    if not power_query:
        return
    add_doc_heading(doc, labels["m_code_heading"], level=2)
    for index, row in enumerate(power_query[:30], 1):
        name = first_value(row, ["TableName", "Name", "QueryName", "table", "name"]) or f"Query {index}"
        expression = first_value(row, ["Expression", "expression", "MExpression"]) or record_text(row)
        add_code_block(doc, str(name), expression, max_chars=2800, labels=labels)


def add_expression_inventory(doc, headers, rows, empty="", labels=None):
    labels = labels or doc_text("pt-BR")
    if not rows:
        add_paragraph(doc, empty)
        return

    inventory_rows = [[row[0], row[1], summarize_expression(row[2], labels)] for row in rows]
    add_records_table(doc, headers + [labels["summary"]], inventory_rows, empty=empty, max_rows=200, labels=labels)

    expressive_rows = [row for row in rows if str(row[2] or "").strip()]
    if not expressive_rows:
        return

    add_doc_heading(doc, labels["technical_expressions"], level=2)
    for table, name, expression in expressive_rows[:80]:
        title = f"{name}"
        if table:
            title = f"{table} - {name}"
        add_code_block(doc, title, expression, max_chars=2200, labels=labels)


def add_code_block(doc, title, code, max_chars=2600, labels=None):
    labels = labels or doc_text("pt-BR")
    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.keep_with_next = True
    title_paragraph.paragraph_format.space_before = Pt(8)
    title_paragraph.paragraph_format.space_after = Pt(3)
    run = title_paragraph.add_run(str(title))
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 109, 133)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="D8DEE8")
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F6F8FA")
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    code_run = paragraph.add_run(normalize_code_text(clip(code, max_chars, labels)))
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(7.5)
    code_run.font.color.rgb = RGBColor(31, 41, 55)
    doc.add_paragraph("")


def build_table_documentation_rows(table_nodes, semantic_tables, schema, labels=None):
    labels = labels or doc_text("pt-BR")
    table_meta = {}
    for row in semantic_tables:
        name = first_value(row, ["Name", "TableName", "Table", "name"])
        if name:
            table_meta[str(name)] = row

    column_counts = {}
    for row in schema:
        table = first_value(row, ["TableName", "Table", "table"])
        column = first_value(row, ["ColumnName", "Name", "Column", "column", "name"])
        if table and column:
            column_counts[str(table)] = column_counts.get(str(table), 0) + 1

    rows = []
    for node in table_nodes:
        name = node.get("label", "")
        meta = table_meta.get(name, {})
        rows.append([
            name,
            column_counts.get(name, ""),
            bool_label(first_value(meta, ["IsHidden", "Hidden", "isHidden"]), labels),
            first_value(meta, ["Description", "description"]) or "",
        ])
    return rows


def build_column_rows(schema, tmschema_columns=None, tmschema_tables=None, labels=None):
    labels = labels or doc_text("pt-BR")
    """Build data dictionary rows merging schema + TMSCHEMA_COLUMNS metadata.

    tmschema_columns uses a numeric TableID foreign key, so we first build a
    TableID → TableName map from tmschema_tables, then index by (table_name, col_name).
    """
    # Build TableID -> TableName map from tmschema_tables
    table_id_to_name = {}
    for row in (tmschema_tables or []):
        tid = first_value(row, ["ID", "TableID", "id"])
        name = first_value(row, ["Name", "TableName", "name"])
        if tid and name:
            table_id_to_name[str(tid)] = str(name)

    # Build lookup (table_name_lower, col_name_lower) -> tmschema row
    tmschema_lookup = {}
    for row in (tmschema_columns or []):
        tid = first_value(row, ["TableID", "Table", "table"])
        col = first_value(row, ["ExplicitName", "Name", "ColumnName", "name"])
        if not col:
            continue
        # Resolve TableID to name if possible
        table_name = table_id_to_name.get(str(tid), "") if tid else ""
        if table_name and not is_system_table(table_name):
            tmschema_lookup[(table_name.lower(), str(col).lower())] = row

    # Normalize DataType codes from TMSCHEMA to readable names
    DATATYPE_MAP = labels["datatype_map"]

    rows = []
    seen = set()
    for row in schema:
        table = first_value(row, ["TableName", "Table", "table"])
        column = first_value(row, ["ColumnName", "Name", "Column", "column", "name"])
        if not table or not column or is_system_table(table):
            continue
        key = (str(table).lower(), str(column).lower())
        if key in seen:
            continue
        seen.add(key)

        # Enrich with tmschema_columns if we resolved the name
        tm = tmschema_lookup.get(key, {})

        raw_type = (
            first_value(tm, ["ExplicitDataType", "DataType", "dataType"])
            or first_value(row, ["DataType", "Type", "type"])
            or ""
        )
        data_type = DATATYPE_MAP.get(str(raw_type).lower(), str(raw_type)) if raw_type else ""

        fmt = (
            first_value(tm, ["FormatString", "Format", "format"])
            or first_value(row, ["FormatString", "Format", "format"])
            or ""
        )

        is_hidden_val = (
            first_value(tm, ["IsHidden", "Hidden", "isHidden"])
            or first_value(row, ["IsHidden", "Hidden", "isHidden"])
        )
        hidden = bool_label(is_hidden_val, labels)

        expression = (
            first_value(tm, ["Expression", "expression"])
            or first_value(row, ["Expression", "expression"])
            or ""
        )

        rows.append([table, column, data_type, fmt, hidden, expression])

    # NOTE: We do NOT add rows directly from tmschema_columns here because its
    # TableID field is a numeric foreign key (e.g. 6107, 15), not a table name.
    # Without a resolved tmschema_tables join, spurious numeric "table names"
    # would appear in the dictionary.
    return rows


def build_expression_rows(raw_records, node_records, kind):
    rows = []
    seen = set()
    for row in raw_records:
        name = first_value(row, ["Name", "MeasureName", "ColumnName", "Measure", "Column", "name"])
        table = first_value(row, ["TableName", "Table", "table"])
        expression = first_value(row, ["Expression", "expression"]) or ""
        if kind == "calc_column" and not expression:
            continue
        # Skip system/internal tables (DateTableTemplate_, LocalDateTable_, etc.)
        if table and is_system_table(table):
            continue
        key = (str(table).lower(), str(name).lower(), str(expression).lower())
        if not name or key in seen:
            continue
        seen.add(key)
        rows.append([table, name, expression])

    if rows:
        return rows

    for node in node_records:
        meta = node.get("meta", {})
        expression = meta.get("expression", "")
        if kind == "calc_column" and not expression:
            continue
        table = meta.get("table", "")
        # Skip system/internal tables from node_records too
        if table and is_system_table(table):
            continue
        rows.append([table, node.get("label", ""), expression])
    return rows


def set_table_width(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def set_table_borders(table, color="D8DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def add_paragraph_bottom_border(paragraph, color="D8DEE8", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_font(cell, color=None, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def bold_cell(cell):
    set_cell_font(cell, bold=True)


def clip(value, limit=900, labels=None):
    labels = labels or doc_text("pt-BR")
    text = doc_value(value)
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 18)].rstrip() + f"\n...[{labels['truncated']}]"


def normalize_code_text(value):
    text = doc_value(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def summarize_m_expression(expression, labels=None):
    labels = labels or doc_text("pt-BR")
    text = doc_value(expression)
    if not text.strip():
        return ""
    steps = re.findall(r"^\s*#?\"?([^\"=\r\n]{2,80})\"?\s*=", text, flags=re.MULTILINE)
    steps = [step.strip() for step in steps if step.strip().lower() not in {"let", "in"}]
    connector_match = re.search(r"([A-Za-z][A-Za-z0-9_]*\.[A-Za-z][A-Za-z0-9_]*)\s*\(", text)
    parts = []
    if connector_match:
        parts.append(f"{labels['connector']}: {connector_match.group(1)}")
    if steps:
        parts.append(f"{labels['steps']}: " + ", ".join(steps[:6]))
    if not parts:
        parts.append(clip(text, 180, labels))
    return " | ".join(parts)


def summarize_expression(expression, labels=None):
    labels = labels or doc_text("pt-BR")
    text = re.sub(r"\s+", " ", doc_value(expression)).strip()
    if not text:
        return ""
    functions = sorted(set(re.findall(r"\b([A-Z][A-Z0-9_]+)\s*\(", text)))
    if functions:
        return f"{labels['functions']}: " + ", ".join(functions[:8])
    return clip(text, 180, labels)


def doc_value(value):
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, default=str)
    return str(value)


def bool_label(value, labels=None):
    labels = labels or doc_text("pt-BR")
    if value in ("", None):
        return ""
    return labels["yes"] if str(value).strip().lower() in {"true", "1", "yes", "sim"} else labels["no"]


def unique_texts(values):
    result = []
    seen = set()
    for value in values:
        text = str(value)
        if text not in seen:
            seen.add(text)
            result.append(text)
    return result


def build_structured_relationships(relationships):
    """
    Normaliza os registros de model.relationships do pbixray para o frontend.

    Campos documentados pela API:
      FromTableName, FromColumnName, ToTableName, ToColumnName, Cardinality, IsActive

    CrossFilteringBehavior pode aparecer como coluna extra dependendo da versão;
    lemos se existir, senão marcamos como desconhecido.
    """
    # Mapa de valores numéricos/texto de Cardinality para símbolo legível
    CARD_MAP = {
        "1": "1:M",  "onetomany": "1:M",  "one_to_many": "1:M",
        "2": "M:1",  "manytoone": "M:1",  "many_to_one": "M:1",
        "3": "1:1",  "onetoone":  "1:1",  "one_to_one":  "1:1",
        "4": "M:M",  "manytomany":"M:M",  "many_to_many":"M:M",
    }
    # CrossFilteringBehavior: 1 = OneDirection / Single, 2 = BothDirections / Both
    CF_MAP = {
        "1": "Single", "onedirection": "Single", "single": "Single",
        "2": "Both", "bothdirections": "Both", "both": "Both",
    }

    result = []
    for rel in relationships:
        from_table  = str(rel.get("FromTableName",  "") or "").strip()
        from_col    = str(rel.get("FromColumnName", "") or "").strip()
        to_table    = str(rel.get("ToTableName",    "") or "").strip()
        to_col      = str(rel.get("ToColumnName",   "") or "").strip()
        cardinality = str(rel.get("Cardinality",    "") or "").strip()
        is_active   = rel.get("IsActive", True)
        cross_raw   = str(rel.get("CrossFilteringBehavior", "") or "").strip()

        if not from_table or not to_table:
            continue

        # Normaliza cardinalidade
        card_key = cardinality.lower().replace(" ", "").replace("-", "")
        card_symbol = CARD_MAP.get(card_key, cardinality if cardinality else "–")

        # Normaliza filtro cruzado
        cf_key = cross_raw.lower().replace(" ", "").replace("-", "")
        cf_label = CF_MAP.get(cf_key, cross_raw if cross_raw else "–")

        # IsActive pode vir como bool, int ou string
        if isinstance(is_active, bool):
            active = is_active
        else:
            active = str(is_active).lower() not in ("false", "0", "no", "none", "")

        result.append({
            "fromTable":  from_table,
            "fromColumn": from_col,
            "toTable":    to_table,
            "toColumn":   to_col,
            "cardinality": card_symbol,
            "crossFilter": cf_label,
            "active":      active,
        })
    return result


def records_from(model, attr, diagnostics):
    try:
        value = getattr(model, attr)
    except Exception as error:
        diagnostics.append(f"{attr}: {error}")
        return []

    if value is None:
        return []
    if hasattr(value, "to_dict"):
        return clean_records(value.to_dict(orient="records"))
    if isinstance(value, list):
        return clean_records(value)
    return []


def list_from(model, attr, diagnostics):
    try:
        value = getattr(model, attr)
    except Exception as error:
        diagnostics.append(f"{attr}: {error}")
        return []
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value]
    return [str(value)]


def clean_records(records):
    cleaned = []
    for record in records:
        if not isinstance(record, dict):
            continue
        cleaned.append({str(key): scalar(value) for key, value in record.items()})
    return cleaned


def record_text(record):
    if not isinstance(record, dict):
        return str(record)
    values = []
    for value in record.values():
        if isinstance(value, (dict, list)):
            values.append(json.dumps(value, ensure_ascii=False, default=str))
        elif value not in ("", None):
            values.append(str(value))
    return " ".join(values)


def scalar(value):
    if value is None:
        return ""
    try:
        if value != value:
            return ""
    except Exception:
        pass
    return value.item() if hasattr(value, "item") else value


def detect_connector_nodes(power_query, datasources):
    detected = {}
    expressions = "\n".join(record_text(row) for row in power_query)

    for connector in CONNECTORS:
        for pattern in connector.get("patterns", []):
            preferred = PREFERRED_PATTERN_CONNECTORS.get(pattern)
            if preferred and connector["name"] != preferred:
                continue
            if has_connector_call(expressions, pattern):
                detected[canonical_connector_name(connector["name"])] = source_node_from_connector(connector, pattern)
                break

        if connector["name"] in detected:
            continue

        for keyword in connector.get("keywords", []):
            if datasource_matches(expressions, keyword):
                detected[canonical_connector_name(connector["name"])] = source_node_from_connector(connector, keyword)
                break

    for source in datasources:
        source_text = " ".join(str(value) for value in source.values() if value)
        matches = []
        for connector in CONNECTORS:
            for keyword in connector.get("keywords", []):
                if not datasource_matches(source_text, keyword):
                    continue
                matches.append((datasource_match_score(source_text, keyword, connector), connector, keyword))
        if matches:
            _, connector, keyword = max(matches, key=lambda item: item[0])
            node = source_node_from_connector(connector, keyword)
            node["meta"]["datasource"] = source
            detected[canonical_connector_name(connector["name"])] = node

    return list(detected.values())


def has_connector_call(text, pattern):
    if not text or not pattern:
        return False
    escaped = re.escape(pattern)
    return re.search(rf"(?<![A-Za-z0-9_]){escaped}\s*\(", text, flags=re.IGNORECASE) is not None


def datasource_matches(text, keyword):
    normalized = str(text).lower()
    keyword = keyword.lower()

    file_extension_keywords = {".xlsx", ".xls", ".xlsb", ".xlsm", ".csv", ".json"}
    if keyword in file_extension_keywords:
        return keyword in normalized

    if keyword in {"sql", "sql server", "microsoft sql", "sql.database", "sql.databases"}:
        sql_server_markers = [
            "sql server",
            "microsoft sql",
            "sql.database",
            "sql.databases",
            "provider=sqlncli",
            "provider=sqloledb",
        ]
        excluded = ["postgresql", "postgre", "mysql", "snowflake", "sqlite"]
        return any(marker in normalized for marker in sql_server_markers) and not any(item in normalized for item in excluded)

    return keyword in normalized


def datasource_match_score(text, keyword, connector):
    normalized = str(text).lower()
    connector_name = connector["name"].lower()
    score = len(keyword)
    if connector_name in normalized:
        score += 100
    if connector_name.replace(" database", "") in normalized:
        score += 35
    if "azure" in normalized and connector_name.startswith("azure"):
        score += 50
    if connector["name"] in {"PostgreSQL", "SQL Server", "Excel", "CSV", "SharePoint"}:
        score += 20
    if connector["name"] == "Excel" and re.search(r"\.xls[xbm]?\b", normalized):
        score += 90
    if connector["name"] == "CSV" and ".csv" in normalized:
        score += 90
    if connector["name"] == "JSON" and ".json" in normalized:
        score += 90
    if connector_name.endswith(" database") and connector_name.replace(" database", "") not in normalized:
        score -= 10
    return score


def query_search_text(query):
    meta = query.get("meta", {})
    return " ".join([
        str(query.get("label", "")),
        str(meta.get("expression", "")),
        str(meta.get("searchText", "")),
        record_text(meta.get("row", {})),
    ])


def source_matches_query(source, query_text):
    pattern = source.get("meta", {}).get("pattern", "")
    label = source.get("label", "")

    if pattern:
        if has_connector_call(query_text, pattern):
            return True
        if datasource_matches(query_text, pattern):
            return True

    if label and datasource_matches(query_text, label):
        return True

    return False


def build_query_nodes(power_query, connectors):
    nodes = []
    for index, row in enumerate(power_query):
        name = first_value(row, ["TableName", "Name", "QueryName", "table", "name"]) or f"Power Query {index + 1}"
        expression = first_value(row, ["Expression", "expression", "MExpression"]) or ""
        full_text = record_text(row)
        connection_path = extract_connection_path_from_m(str(expression))
        nodes.append({
            "id": f"query:{slug(name)}",
            "type": "query",
            "label": str(name),
            "icon": "M",
            "meta": {
                "expression": str(expression),
                "searchText": full_text,
                "row": row,
                "connectionPath": connection_path,
            },
        })

    if not nodes and connectors:
        for connector in connectors:
            nodes.append({
                "id": f"query:{slug(connector['label'])}-query",
                "type": "query",
                "label": f"{connector['label']} Query",
                "icon": "M",
                "meta": {"expression": connector["meta"].get("pattern", ""), "connectionPath": ""},
            })

    return nodes


def extract_connection_path_from_m(expression: str) -> str:
    """
    Extracts the connection arguments from the Source = ConnectorFunction(...) step.
    Returns a human-readable string like 'server.database.windows.net › SalesDB'
    or a file path / URL.
    """
    if not expression:
        return ""

    # Connector function patterns (e.g. Sql.Database("server", "db"))
    connector_call = re.search(
        r'(?:Source\s*=\s*)?([A-Za-z][A-Za-z0-9_]*\.[A-Za-z][A-Za-z0-9_]*)\s*\(([^)]{0,600})\)',
        expression
    )
    if connector_call:
        args_str = connector_call.group(2)
        args = re.findall(r'"([^"]{1,300})"', args_str)
        if args:
            return " › ".join(args)

    # SharePoint / web URL
    url_match = re.search(r'https?://[^\s"\']{4,300}', expression)
    if url_match:
        return url_match.group(0)

    # Windows file path
    path_match = re.search(r'[A-Za-z]:[/\\][^"\'<>\r\n]{3,200}', expression)
    if path_match:
        return path_match.group(0)

    # UNC path
    unc_match = re.search(r'\\\\[^"\'<>\r\n]{3,200}', expression)
    if unc_match:
        return unc_match.group(0)

    # First quoted string
    first_quoted = re.search(r'"([^"]{3,200})"', expression)
    if first_quoted:
        return first_quoted.group(1)

    return ""


def build_table_nodes(tables, schema, semantic_tables, query_nodes):
    names = set()
    hidden = set()

    for row in semantic_tables:
        table = first_value(row, ["Name", "TableName", "Table", "name"])
        if not table:
            continue
        is_hidden = str(first_value(row, ["IsHidden", "Hidden", "isHidden"])).lower() in ("true", "1", "yes")
        if is_hidden:
            hidden.add(str(table))
        elif not is_system_table(table):
            names.add(str(table))

    if not names:
        names = {str(table) for table in tables if table and not is_system_table(table)}

    for row in schema:
        table = first_value(row, ["TableName", "Table", "table"])
        if table and table not in hidden and not is_system_table(table):
            names.add(str(table))

    query_names = {node["label"] for node in query_nodes if not is_system_table(node["label"])}
    if query_names:
        names = {name for name in names if name in query_names or not is_system_table(name)}

    if not names:
        names = {node["label"] for node in query_nodes}

    return [{"id": f"model:{slug(name)}", "type": "model", "label": name, "icon": "TBL", "meta": {}} for name in sorted(names)]


def is_system_table(name):
    normalized = str(name).strip().lower()
    return (
        not normalized
        or normalized.startswith("<")
        or normalized.startswith("datetabletemplate")
        or normalized.startswith("localdatetable")
        or normalized.startswith("date table template")
        or normalized.startswith("_")
    )


def build_measure_nodes(measures):
    nodes = []
    for index, row in enumerate(measures):
        name = first_value(row, ["Name", "MeasureName", "Measure", "name"]) or f"Measure {index + 1}"
        table = first_value(row, ["TableName", "Table", "table"])
        expression = first_value(row, ["Expression", "expression"]) or ""
        nodes.append({
            "id": f"measure:{slug(table)}:{slug(name)}" if table else f"measure:{slug(name)}",
            "type": "measure",
            "label": str(name),
            "icon": "DAX",
            "meta": {"table": str(table) if table else "", "expression": str(expression), "row": row},
        })
    return nodes


def build_measure_dependency_edges(measure, measure_nodes, table_nodes):
    edges = []
    expression = measure["meta"].get("expression", "")
    if not expression:
        return edges

    expression_lower = expression.lower()
    for table in table_nodes:
        label = table["label"]
        if f"'{label.lower()}'" in expression_lower or re.search(rf"\b{re.escape(label.lower())}\b\s*\[", expression_lower):
            edges.append(edge(table["id"], measure["id"], "referenced in DAX"))

    for other in measure_nodes:
        if other["id"] == measure["id"]:
            continue
        if f"[{other['label'].lower()}]" in expression_lower:
            edges.append(edge(other["id"], measure["id"], "measure dependency"))

    return edges


def build_calc_column_dependency_edges(calc_col, measure_nodes, table_nodes):
    edges = []
    expression = calc_col["meta"].get("expression", "")
    if not expression:
        return edges
    expression_lower = expression.lower()
    for other in measure_nodes:
        if f"[{other['label'].lower()}]" in expression_lower:
            edges.append(edge(other["id"], calc_col["id"], "referenced in calc column"))
    return edges


def build_calc_column_nodes(raw_columns):
    """Build calc_column nodes from dax_columns records.

    pbixray's dax_columns table contains all model columns; we keep only
    those that have an Expression (i.e. are DAX-calculated columns).

    Power BI stores calculated columns once per user table AND once again
    inside internal hierarchy/date-template tables (LocalDateTable_*,
    DateTableTemplate_*, etc). We deduplicate in two steps:
      1. Skip rows whose TableName is a system/internal table.
      2. Deduplicate by (slug_name, normalised_expression) to catch the same
         column surfaced under different table name spellings that slipped
         past is_system_table.
    """
    nodes = []
    seen_id = set()
    seen_expr_key = set()
    for index, row in enumerate(raw_columns):
        expression = first_value(row, ["Expression", "expression"]) or ""
        if not expression.strip():
            continue  # skip plain (non-calculated) columns
        name = first_value(row, ["Name", "ColumnName", "Column", "name"]) or f"Calc Column {index + 1}"
        table = first_value(row, ["TableName", "Table", "table"]) or ""

        # Skip columns belonging to Power BI internal/system tables
        if is_system_table(table):
            continue

        node_id = f"calc_column:{slug(table)}:{slug(name)}" if table else f"calc_column:{slug(name)}"
        if node_id in seen_id:
            continue
        seen_id.add(node_id)

        # Secondary dedup: same column name + same expression avoids phantom
        # duplicates when pbixray surfaces the column under a non-system table
        # alias with a different name (e.g. auto-date hierarchy sub-tables).
        expr_key = (slug(str(name)), re.sub(r"\s+", " ", str(expression).strip().lower()))
        if expr_key in seen_expr_key:
            continue
        seen_expr_key.add(expr_key)

        nodes.append({
            "id": node_id,
            "type": "calc_column",
            "label": str(name),
            "icon": "CC",
            "meta": {"table": str(table), "expression": str(expression), "row": row},
        })
    return nodes


def extract_pages_from_layout(path: Path):
    """Extract report pages (sections) from the PBIX layout file.

    Returns a list of dicts:
      {
        "name":         str,   # displayName shown in Power BI
        "ordinal":      int,   # page order (0-based)
        "visualCount":  int,   # number of visual containers on the page
        "width":        int,   # canvas width in pixels
        "height":       int,   # canvas height in pixels
      }
    """
    try:
        with zipfile.ZipFile(path) as archive:
            layout_name = next((n for n in archive.namelist() if n.lower() == "report/layout"), "")
            if not layout_name:
                return []
            raw = archive.read(layout_name)
    except Exception:
        return []

    text = raw.decode("utf-16le", errors="ignore")
    if text.count("{") < 5:
        text = raw.decode("utf-8", errors="ignore")

    try:
        layout = json.loads(text)
    except Exception:
        return []

    pages = []
    for section in layout.get("sections", []):
        display_name = section.get("displayName") or section.get("name") or f"Page {len(pages) + 1}"
        ordinal      = int(section.get("ordinal", len(pages)))
        containers   = section.get("visualContainers", [])
        visual_count = len(containers)
        width        = int(section.get("width",  1280))
        height       = int(section.get("height",  720))
        pages.append({
            "name":        str(display_name),
            "ordinal":     ordinal,
            "visualCount": visual_count,
            "width":       width,
            "height":      height,
        })

    pages.sort(key=lambda p: p["ordinal"])
    return pages


def extract_visuals_from_layout(path: Path):
    try:
        with zipfile.ZipFile(path) as archive:
            layout_name = next((name for name in archive.namelist() if name.lower() == "report/layout"), "")
            if not layout_name:
                return []
            raw = archive.read(layout_name)
    except Exception:
        return []

    text = raw.decode("utf-16le", errors="ignore")
    if text.count("{") < 5:
        text = raw.decode("utf-8", errors="ignore")

    visuals = []
    try:
        layout = json.loads(text)
        for section in layout.get("sections", []):
            for index, container in enumerate(section.get("visualContainers", [])):
                config = parse_json_string(container.get("config"))
                query = parse_json_string(container.get("query"))
                visual_type = nested_value(config, ["singleVisual", "visualType"]) or "Visual"
                title = visual_title(config) or f"{titleize(visual_type)} {index + 1}"
                refs = sorted(extract_query_refs(config) | extract_query_refs(query))
                visuals.append({"name": title, "type": titleize(visual_type), "refs": refs})
    except Exception:
        for index, match in enumerate(re.finditer(r'"visualType"\s*:\s*"([^"]+)"', text)):
            visual_type = titleize(match.group(1))
            start = max(0, match.start() - 3000)
            end = min(len(text), match.end() + 3000)
            refs = sorted(extract_refs_from_text(text[start:end]))
            visuals.append({"name": f"{visual_type} {index + 1}", "type": visual_type, "refs": refs})

    deduped = {}
    for visual in visuals:
        key = visual["name"]
        if key in deduped:
            deduped[key]["refs"] = sorted(set(deduped[key]["refs"]) | set(visual["refs"]))
        else:
            deduped[key] = visual
    return list(deduped.values())[:24]


def build_visual_nodes(visuals):
    return [
        {
            "id": f"visual:{slug(visual['name'])}",
            "type": "visual",
            "label": visual["name"],
            "icon": "VIS",
            "meta": {"visualType": visual.get("type", ""), "refs": visual.get("refs", [])},
        }
        for visual in visuals
    ]


def build_visual_edges(visual_nodes, measure_nodes, table_nodes, calc_column_nodes=None):
    edges = []
    measures_by_label = {node["label"].lower(): node for node in measure_nodes}
    tables_by_label = {node["label"].lower(): node for node in table_nodes}

    # Index calc columns by bare name AND by "table.columnname" composite so that
    # refs extracted from the layout (which often appear as "Entity"/"Property" pairs
    # or as a dotted "queryRef") can match either way.
    calc_cols_by_label = {}
    calc_cols_by_composite = {}
    for node in (calc_column_nodes or []):
        bare = node["label"].lower()
        calc_cols_by_label[bare] = node
        table = node["meta"].get("table", "")
        if table:
            composite = f"{table.lower()}.{bare}"
            calc_cols_by_composite[composite] = node

    for visual in visual_nodes:
        refs = {str(ref).lower() for ref in visual["meta"].get("refs", [])}
        matched = False

        for label, measure in measures_by_label.items():
            if label in refs:
                edges.append(edge(measure["id"], visual["id"], "used in visual"))
                matched = True

        for label, calc_col in calc_cols_by_label.items():
            if label in refs:
                edges.append(edge(calc_col["id"], visual["id"], "used in visual"))
                matched = True

        # Secondary pass: try composite "table.column" match for calc columns that
        # weren't caught by the bare name (e.g. when the layout stores them as
        # "Entity":"Sales","Property":"YTD Category" and the ref set contains
        # "sales.ytd category" after joining).
        for composite, calc_col in calc_cols_by_composite.items():
            if composite in refs and calc_col["id"] not in {e["from"] for e in edges if e["to"] == visual["id"]}:
                edges.append(edge(calc_col["id"], visual["id"], "used in visual"))
                matched = True

        for label, table in tables_by_label.items():
            if label in refs:
                edges.append(edge(table["id"], visual["id"], "used in visual"))
                matched = True

        if not matched and len(measure_nodes) == 1:
            edges.append(edge(measure_nodes[0]["id"], visual["id"], "used in visual"))
        elif not matched and not measure_nodes and len(table_nodes) == 1:
            edges.append(edge(table_nodes[0]["id"], visual["id"], "used in visual"))

    return edges


def parse_json_string(value):
    if isinstance(value, dict):
        return value
    if not isinstance(value, str) or not value.strip():
        return {}
    try:
        return json.loads(value)
    except Exception:
        return {}


def nested_value(data, path):
    current = data
    for key in path:
        if not isinstance(current, dict):
            return ""
        current = current.get(key)
    return current if isinstance(current, str) else ""


def visual_title(config):
    text = json.dumps(config, ensure_ascii=False)
    match = re.search(r'"titleText"\s*:\s*"([^"]+)"', text)
    return match.group(1) if match else ""


def extract_query_refs(value):
    text = json.dumps(value, ensure_ascii=False) if not isinstance(value, str) else value
    return extract_refs_from_text(text)


def extract_refs_from_text(text):
    refs = set()
    for pattern in [
        r'"queryRef"\s*:\s*"([^"]+)"',
        r'"Entity"\s*:\s*"([^"]+)"',
        r'"Property"\s*:\s*"([^"]+)"',
        r'"Measure"\s*:\s*"([^"]+)"',
        r'"Column"\s*:\s*"([^"]+)"',
        r'"Name"\s*:\s*"([^"]+)"',
    ]:
        for match in re.finditer(pattern, text):
            value = match.group(1)
            # Keep the full dotted path (e.g. "Sales.YTD Category") so that
            # build_visual_edges can attempt a "table.column" composite match.
            refs.add(value.lower())
            # Also split into individual parts so bare column/measure names match too.
            for part in re.split(r"[.\[\]]+", value):
                clean = part.strip("'\" ")
                if clean:
                    refs.add(clean.lower())
    return refs


def source_node_from_connector(connector, matched_by):
    display_name = canonical_connector_name(connector["name"])
    return source_node(
        display_name,
        connector.get("icon", initials(connector["name"])),
        matched_by,
        icon_url=connector.get("iconUrl", ""),
        doc=connector.get("doc", ""),
        image=connector.get("image", ""),
    )


def canonical_connector_name(name):
    return CANONICAL_CONNECTOR_NAMES.get(name, name)


def source_node(name, icon, pattern, icon_url="", doc="", image=""):
    return {
        "id": f"source:{slug(name)}",
        "type": "source",
        "label": name,
        "icon": icon,
        "iconUrl": icon_url,
        "meta": {"pattern": pattern, "doc": doc, "image": image},
    }


def edge(source, target, label):
    return {"id": f"{source}->{target}", "from": source, "to": target, "label": label}


def first_value(record, keys):
    for key in keys:
        if key in record and record[key] not in ("", None):
            return record[key]
    return ""


def slug(value):
    value = str(value).lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "item"


def initials(value):
    parts = [part for part in re.split(r"[^A-Za-z0-9]+", str(value)) if part]
    return ("".join(part[0] for part in parts[:3]).upper() or "SRC")[:4]


def titleize(value):
    value = re.sub(r"[_-]+", " ", str(value))
    value = re.sub(r"([a-z])([A-Z])", r"\1 \2", value)
    return value.title()


def unique_by_id(items):
    result = {}
    for item in items:
        result[item["id"]] = item
    return list(result.values())


def unique_edges(items):
    result = {}
    for item in items:
        result[item["id"]] = item
    return list(result.values())


def is_port_open(port):
    try:
        with socket.create_connection((HOST, port), timeout=0.5):
            return True
    except OSError:
        return False


def is_healthy(port):
    try:
        with urlopen(f"http://{HOST}:{port}/api/health", timeout=0.8) as response:
            return response.status == 200
    except Exception:
        return False


def find_port(start_port):
    for port in range(start_port, start_port + 20):
        if not is_port_open(port):
            return port
    raise RuntimeError("Nenhuma porta local livre encontrada entre 4173 e 4192.")


def requested_port():
    for index, arg in enumerate(sys.argv):
        if arg == "--port" and index + 1 < len(sys.argv):
            return int(sys.argv[index + 1])
        if arg.startswith("--port="):
            return int(arg.split("=", 1)[1])
    return DEFAULT_PORT


def open_browser(port):
    webbrowser.open(f"http://{HOST}:{port}")


def write_port_file(port):
    PORT_FILE.write_text(f"http://{HOST}:{port}", encoding="utf-8")


def main():
    should_open = "--open" in sys.argv
    start_port = requested_port()

    if is_healthy(start_port):
        write_port_file(start_port)
        if should_open:
            open_browser(start_port)
        print(f"BI Flow Mapper already running at http://{HOST}:{start_port}")
        return

    port = find_port(start_port)
    server = ThreadingHTTPServer((HOST, port), Handler)
    write_port_file(port)

    if should_open:
        threading.Timer(0.4, open_browser, args=(port,)).start()

    print(f"BI Flow Mapper running at http://{HOST}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
